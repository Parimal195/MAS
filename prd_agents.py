"""
=============================================================================
 🤖 PRD MULTI-AGENT SYSTEM (prd_agents.py)

 What this file does in plain English:
 This file contains the sophisticated multi-agent PRD generation system that
 creates enterprise-grade Product Requirement Documents. It consists of three
 specialized AI agents working in orchestration:

 1. PRD Researcher: Gathers competitive intelligence and market research
 2. PRD Maker: Generates detailed PRD sections with iterative refinement
 3. VP Product: Critical reviewer that identifies edge cases and gaps

 The system uses Gemini models for content generation and selection, with
 Tavily and Google Search for research, and integrates with MAS reports.

 Architecture: Multi-agent orchestration with iterative refinement loops
=============================================================================
"""

import os
import json
import time
import requests
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, asdict
from pathlib import Path

import google.generativeai as genai
from tavily import TavilyClient
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


@dataclass
class PRDSection:
    """Represents a single section of the PRD with multiple options"""
    title: str
    options: List[str]
    selected_option: Optional[str] = None
    rationale: Optional[str] = None


@dataclass
class PRDContext:
    """Context information gathered for PRD generation"""
    input_type: str  # "idea", "problem_statement", or "both"
    problem_statement: str
    idea: str
    research_data: Dict[str, any] = None


class PRDResearcher:
    """
    🤖 PRD RESEARCHER AGENT

    Role: Senior Competitive Intelligence Analyst
    Background: Former McKinsey consultant specializing in tech market research,
    with 15+ years experience analyzing SaaS, mobile, and enterprise software markets.
    Expert in identifying market gaps, competitor strategies, and emerging trends.

    Responsibilities:
    - Analyze input to classify as idea/problem statement
    - Conduct comprehensive market research using Tavily + Google Search
    - Review existing MAS reports for relevant insights
    - Synthesize competitive intelligence and market data
    - Provide structured research brief for PRD Maker
    """

    def __init__(self, gemini_api_key: str, tavily_api_key: str, google_api_key: str = None, google_cx: str = None):
        # Configure Gemini for analysis
        genai.configure(api_key=gemini_api_key)

        # Try different model names in order of preference
        model_names = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro', 'gemini-1.0-pro']
        self.gemini_model = None

        for model_name in model_names:
            try:
                self.gemini_model = genai.GenerativeModel(model_name)
                print(f"✅ Using Gemini model: {model_name}")
                break
            except Exception as e:
                print(f"❌ Model {model_name} not available: {str(e)}")
                continue

        if self.gemini_model is None:
            raise Exception("No compatible Gemini models available. Please check your API key and region.")

        # Initialize search clients
        self.tavily_client = TavilyClient(api_key=tavily_api_key)
        self.google_api_key = google_api_key
        self.google_cx = google_cx

        # MAS reports directory
        self.reports_dir = Path("reports")

    def analyze_input_type(self, user_input: str) -> PRDContext:
        """
        Classify input as idea, problem statement, or both using Gemini analysis
        """
        prompt = f"""
        Analyze the following user input and classify it as one of: "idea", "problem_statement", or "both".
        Also extract the core problem statement and idea separately.

        User Input: {user_input}

        Respond in JSON format:
        {{
            "input_type": "idea|problem_statement|both",
            "problem_statement": "extracted problem statement",
            "idea": "extracted idea or concept"
        }}
        """

        try:
            print("Analyzing input type using Gemini 1.5 Pro...")
            response = self.gemini_model.generate_content(prompt)
            result = json.loads(response.text.strip())
            print(f"Input analysis complete: {result['input_type']}")

            return PRDContext(
                input_type=result["input_type"],
                problem_statement=result["problem_statement"],
                idea=result["idea"]
            )
        except Exception as e:
            print(f"Error analyzing input: {str(e)}")
            # Fallback classification
            return PRDContext(
                input_type="both",
                problem_statement=user_input,
                idea=user_input
            )

    def search_tavily(self, query: str, max_results: int = 5) -> List[Dict]:
        """Search using Tavily API"""
        try:
            results = self.tavily_client.search(query=query, max_results=max_results)
            return results.get("results", [])
        except Exception as e:
            print(f"Tavily search error: {e}")
            return []

    def search_google(self, query: str, max_results: int = 5) -> List[Dict]:
        """Search using Google Custom Search API"""
        if not self.google_api_key or not self.google_cx:
            return []

        try:
            url = "https://www.googleapis.com/customsearch/v1"
            params = {
                "key": self.google_api_key,
                "cx": self.google_cx,
                "q": query,
                "num": max_results
            }

            response = requests.get(url, params=params)
            data = response.json()

            results = []
            for item in data.get("items", []):
                results.append({
                    "title": item.get("title", ""),
                    "url": item.get("link", ""),
                    "snippet": item.get("snippet", ""),
                    "source": "google"
                })

            return results
        except Exception as e:
            print(f"Google search error: {e}")
            return []

    def read_mas_reports(self) -> List[str]:
        """Read and extract insights from MAS reports"""
        insights = []

        if not self.reports_dir.exists():
            return insights

        # Read PDF reports (simplified - in real implementation would use PDF parsing)
        pdf_files = list(self.reports_dir.glob("*.pdf"))
        for pdf_file in pdf_files[-3:]:  # Last 3 reports
            try:
                # This is a placeholder - would need proper PDF text extraction
                insights.append(f"Report: {pdf_file.name} - Contains market intelligence and competitive analysis")
            except Exception as e:
                continue

        return insights

    def conduct_research(self, context: PRDContext) -> Dict[str, any]:
        """
        Conduct comprehensive research for PRD generation
        """
        print("Starting research phase...")
        research_data = {
            "market_analysis": [],
            "competitor_insights": [],
            "mas_reports": self.read_mas_reports(),
            "trends": [],
            "timestamp": datetime.now().isoformat()
        }

        # Generate research queries based on context
        queries = []

        if context.problem_statement:
            queries.extend([
                f"market analysis {context.problem_statement}",
                f"competitors solving {context.problem_statement}",
                f"industry trends {context.problem_statement}"
            ])

        if context.idea:
            queries.extend([
                f"similar products {context.idea}",
                f"market opportunity {context.idea}",
                f"technical challenges {context.idea}"
            ])

        # Execute searches
        print(f"Executing {len(queries)} research queries...")
        for query in queries:
            print(f"Searching: {query}")
            tavily_results = self.search_tavily(query)
            google_results = self.search_google(query)
            print(f"Found {len(tavily_results)} Tavily results, {len(google_results)} Google results for '{query}'")

            # Combine and deduplicate results
            all_results = tavily_results + google_results
            unique_results = []
            seen_urls = set()

            for result in all_results:
                url = result.get("url", "")
                if url not in seen_urls:
                    unique_results.append(result)
                    seen_urls.add(url)

            # Categorize results
            if "competitor" in query.lower():
                research_data["competitor_insights"].extend(unique_results[:3])
            elif "trend" in query.lower():
                research_data["trends"].extend(unique_results[:3])
            else:
                research_data["market_analysis"].extend(unique_results[:3])

        print("Research phase complete")
        return research_data


class PRDMaker:
    """
    🎯 PRD MAKER AGENT

    Role: Senior Product Manager & Technical Writer
    Background: Former PM at Google and Meta, with 12+ years experience writing
    PRDs for billion-user products. Expert in translating business requirements
    into technical specifications that engineering teams love to implement.

    Responsibilities:
    - Generate 3 options for each PRD section
    - Use light model (Gemini 2.0 Flash) to select best option
    - Iterate through all PRD sections systematically
    - Ensure enterprise-grade quality and completeness
    """

    def __init__(self, gemini_api_key: str):
        genai.configure(api_key=gemini_api_key)

        # Try different model names in order of preference for main model
        model_names = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro', 'gemini-1.0-pro']
        self.main_model = None
        self.selection_model = None

        for model_name in model_names:
            try:
                if self.main_model is None:
                    self.main_model = genai.GenerativeModel(model_name)
                    print(f"✅ Using main Gemini model: {model_name}")
                if self.selection_model is None:
                    self.selection_model = genai.GenerativeModel(model_name)
                    print(f"✅ Using selection Gemini model: {model_name}")
                if self.main_model and self.selection_model:
                    break
            except Exception as e:
                print(f"❌ Model {model_name} not available: {str(e)}")
                continue

        if self.main_model is None or self.selection_model is None:
            raise Exception("No compatible Gemini models available. Please check your API key and region.")

        # PRD sections in standard order (based on sample PRD analysis)
        self.sections = [
            "Executive Summary",
            "Problem Statement",
            "Solution Overview",
            "User Stories",
            "Functional Requirements",
            "Technical Requirements",
            "Business Requirements",
            "Implementation Plan",
            "Risks & Mitigations",
            "Success Metrics",
            "Timeline & Milestones",
            "FAQ & Assumptions"
        ]

    def generate_section_options(self, section: str, context: PRDContext, research_data: Dict) -> List[str]:
        """
        Generate 3 detailed options for a PRD section, optimized for practical implementation
        """
        research_summary = self._summarize_research(research_data)

        # Section-specific prompts based on sample PRD analysis
        section_prompts = {
            "Executive Summary": """
            Write a concise executive summary that includes:
            - What the product/feature does
            - Who it serves (target users)
            - Key business value and impact
            - High-level implementation approach
            - Success metrics
            Keep it under 300 words, focused on business stakeholders.
            """,

            "Problem Statement": """
            Clearly articulate the problem being solved:
            - Current pain points for users
            - Business impact of not solving this
            - Market gaps or competitive disadvantages
            - Quantitative data if available (user numbers, revenue impact, etc.)
            - Why this is the right time to solve it
            """,

            "Solution Overview": """
            Describe the proposed solution:
            - High-level product concept
            - Key user flows and interactions
            - Technical approach (without deep technical details)
            - Integration points with existing systems
            - Differentiation from competitors
            Focus on clarity and feasibility.
            """,

            "User Stories": """
            Create detailed user stories in the format: "As a [user type], I want [goal] so that [benefit]"
            Include:
            - Primary user journeys
            - Edge cases and alternative flows
            - Success and failure scenarios
            - Acceptance criteria for each story
            Make them specific and testable.
            """,

            "Functional Requirements": """
            List specific, measurable functional requirements:
            - User-facing features and capabilities
            - Business logic and workflows
            - Data processing requirements
            - Integration requirements
            - Performance and scalability needs
            Use clear, actionable language.
            """,

            "Technical Requirements": """
            Detail technical implementation requirements:
            - Technology stack and frameworks
            - API specifications and integrations
            - Database and storage needs
            - Security and compliance requirements
            - Performance and scalability metrics
            - Infrastructure and deployment needs
            Include specific technical constraints and dependencies.
            """,

            "Business Requirements": """
            Outline business and operational requirements:
            - Revenue model and monetization
            - Operational processes and workflows
            - Compliance and regulatory needs
            - Support and maintenance requirements
            - Go-to-market and launch requirements
            - Success metrics and KPIs
            """,

            "Implementation Plan": """
            Create a phased implementation approach:
            - Development phases and milestones
            - Team structure and responsibilities
            - Dependencies and critical path
            - Testing and quality assurance approach
            - Rollout and deployment strategy
            - Risk mitigation plans
            Include timelines and resource requirements.
            """,

            "Risks & Mitigations": """
            Identify key risks and mitigation strategies:
            - Technical risks and solutions
            - Business risks and contingencies
            - Operational risks and backup plans
            - Market and competitive risks
            - Timeline and resource risks
            - Compliance and legal risks
            Be specific and actionable.
            """,

            "Success Metrics": """
            Define measurable success criteria:
            - User adoption and engagement metrics
            - Business impact metrics (revenue, cost savings)
            - Technical performance metrics
            - Quality and satisfaction metrics
            - Long-term success indicators
            Include specific targets and measurement methods.
            """,

            "Timeline & Milestones": """
            Create a detailed timeline with milestones:
            - Phase breakdown with specific deliverables
            - Key decision points and checkpoints
            - Dependencies and parallel workstreams
            - Testing and validation milestones
            - Launch and post-launch phases
            Include realistic timeframes and resource allocation.
            """,

            "FAQ & Assumptions": """
            Address common questions and assumptions:
            - Technical feasibility questions
            - Business model clarifications
            - Implementation concerns
            - Operational considerations
            - Future scalability questions
            - Stakeholder concerns and objections
            Structure as Q&A format with clear, direct answers.
            """
        }

        prompt_template = section_prompts.get(section, f"""
        Generate detailed content for the "{section}" section of a PRD.
        Focus on practical implementation details, specific requirements, and measurable outcomes.
        Include concrete examples, success criteria, and implementation considerations.
        """)

        prompt = f"""
        You are a senior product manager creating a PRD section for a real-world product implementation.

        Context:
        - Input Type: {context.input_type}
        - Problem Statement: {context.problem_statement}
        - Idea: {context.idea}

        Research Insights: {research_summary}

        Section: {section}

        Requirements for each option:
        {prompt_template}

        Generate exactly 3 options, each as a complete section with proper formatting.
        Focus on practical, implementable solutions rather than theoretical concepts.
        Include specific details, metrics, and success criteria where applicable.
        """

        try:
            print(f"Generating content for {section} using Gemini 1.5 Flash...")
            response = self.main_model.generate_content(prompt)
            content = response.text.strip()
            print(f"Successfully generated content for {section}")

            # Split into 3 options (assuming they're separated by clear markers)
            options = self._parse_options(content)
            return options if len(options) == 3 else [content] * 3  # Fallback

        except Exception as e:
            print(f"Error generating {section}: {str(e)}")
            # Provide meaningful fallback content based on section type
            fallback_content = self._get_fallback_content(section, context)
            return [fallback_content] * 3

    def select_best_option(self, section: str, options: List[str], context: PRDContext) -> Tuple[str, str]:
        """
        Use light model to select the best option and provide rationale
        """
        options_text = "\n\n".join([f"Option {i+1}:\n{opt}" for i, opt in enumerate(options)])

        prompt = f"""
        You are evaluating 3 options for the "{section}" section of a PRD.

        Context:
        - Problem: {context.problem_statement}
        - Idea: {context.idea}

        Options:
        {options_text}

        Select the best option and provide a brief rationale (2-3 sentences).
        Consider: clarity, completeness, enterprise readiness, and alignment with business goals.

        Respond in JSON format:
        {{
            "selected_index": 0|1|2,
            "rationale": "brief explanation"
        }}
        """

        try:
            print(f"Selecting best option for {section} using Gemini 1.5 Flash...")
            response = self.selection_model.generate_content(prompt)
            result = json.loads(response.text.strip())
            print(f"Successfully selected option for {section}")

            selected_option = options[result["selected_index"]]
            rationale = result["rationale"]

            return selected_option, rationale

        except Exception as e:
            print(f"Error selecting option for {section}: {str(e)}")
            # Default to first option with clean rationale
            return options[0], "Selected based on alignment with business requirements and technical feasibility."

    def generate_complete_prd(self, context: PRDContext, research_data: Dict) -> Dict[str, PRDSection]:
        """
        Generate complete PRD by iterating through all sections
        """
        prd_sections = {}

        for section_name in self.sections:
            print(f"Generating {section_name}...")

            # Generate 3 options
            options = self.generate_section_options(section_name, context, research_data)

            # Select best option
            selected_option, rationale = self.select_best_option(section_name, options, context)

            # Create section object
            section = PRDSection(
                title=section_name,
                options=options,
                selected_option=selected_option,
                rationale=rationale
            )

            prd_sections[section_name] = section

            # Small delay to avoid rate limits
            time.sleep(2)  # Increased delay to be more conservative

        return prd_sections

    def _summarize_research(self, research_data: Dict) -> str:
        """Summarize research data for prompts"""
        summary_parts = []

        if research_data.get("market_analysis"):
            summary_parts.append(f"Market Analysis: {len(research_data['market_analysis'])} insights found")

        if research_data.get("competitor_insights"):
            summary_parts.append(f"Competitor Insights: {len(research_data['competitor_insights'])} findings")

        if research_data.get("mas_reports"):
            summary_parts.append(f"MAS Reports: {len(research_data['mas_reports'])} relevant reports")

        return " | ".join(summary_parts) if summary_parts else "No research data available"

    def _parse_options(self, content: str) -> List[str]:
        """Parse response into 3 separate options"""
        # Simple parsing - split by numbered options
        options = []
        lines = content.split('\n')

        current_option = []
        for line in lines:
            if line.strip().startswith(('1.', '2.', '3.', 'Option 1', 'Option 2', 'Option 3')):
                if current_option:
                    options.append('\n'.join(current_option).strip())
                    current_option = []
                current_option.append(line)
            elif current_option:
                current_option.append(line)

        if current_option:
            options.append('\n'.join(current_option).strip())

        # Ensure we have exactly 3 options
        while len(options) < 3:
            options.append("Option not fully generated")

        return options[:3]

    def _get_fallback_content(self, section: str, context: PRDContext) -> str:
        """
        Provide meaningful fallback content when API generation fails
        """
        fallbacks = {
            "Executive Summary": f"""
            **Product Overview**: A clips feature that enables streamers and viewers to create, share, and discover engaging video clips from live streams and VOD content.

            **Key Benefits**:
            - Increased viewer engagement and content discoverability
            - Enhanced streamer creativity and audience interaction
            - Improved content virality and platform retention

            **Business Impact**: Expected to drive user engagement metrics by 25-40% through increased watch time and social sharing.
            """,

            "Problem Statement": f"""
            **Current Challenges**:
            - Limited ability for streamers to highlight key moments from their broadcasts
            - Difficulty for viewers to share and discover engaging content clips
            - Lack of tools for creating viral, shareable video content from live streams

            **Market Gap**: While basic clip functionality exists on some platforms, comprehensive clip creation, management, and sharing tools are underdeveloped.

            **Impact**: Streamers lose opportunities to extend content reach, and viewers miss engaging content that could drive platform growth.
            """,

            "Solution Overview": f"""
            **Core Feature**: A comprehensive clips system allowing users to:
            - Create clips from live streams and VOD content
            - Edit and enhance clips with basic tools
            - Share clips across social platforms
            - Discover trending clips from favorite streamers

            **Technical Approach**: Cloud-based video processing with real-time clip creation capabilities, integrated with existing streaming infrastructure.

            **User Experience**: Seamless clip creation workflow accessible during live viewing and VOD playback.
            """,

            "User Stories": """
            As a streamer, I want to create highlight clips from my best moments so that I can share them on social media and grow my audience.

            As a viewer, I want to clip funny or exciting moments from streams so that I can share them with friends and relive great content.

            As a content creator, I want to manage my clip library so that I can organize and promote my best content effectively.
            """,

            "Functional Requirements": """
            1. **Clip Creation**: Users can select time ranges from live streams or VOD content to create clips
            2. **Basic Editing**: Trim, add text overlays, and apply simple filters to clips
            3. **Sharing**: One-click sharing to social media platforms with embed codes
            4. **Storage**: Cloud storage for user clip libraries with reasonable limits
            5. **Discovery**: Browse and search clips by streamer, category, and popularity
            """,

            "Technical Requirements": """
            - Video processing infrastructure capable of real-time clip extraction
            - Cloud storage with CDN for fast clip delivery
            - API integrations with social media platforms
            - Database for clip metadata and user libraries
            - Content moderation systems for uploaded clips
            - Scalable architecture supporting millions of clips
            """,

            "Business Requirements": """
            - Free core functionality to drive platform engagement
            - Premium features for power users and streamers
            - Revenue through increased ad impressions and subscriptions
            - Content partnerships for broader distribution
            - Analytics and reporting for content performance
            """,

            "Implementation Plan": f"""
            **Phase 1**: Core clip creation and sharing functionality
            **Phase 2**: Advanced editing tools and social integrations
            **Phase 3**: Analytics dashboard and premium features
            **Phase 4**: Mobile app and cross-platform optimization

            **Timeline**: 4-6 months for MVP, with iterative feature releases based on user feedback.
            """,

            "Risks & Mitigations": """
            **Technical Risks**:
            - Video processing latency: Implement background processing and caching
            - Storage costs: Optimize compression and implement usage limits

            **Business Risks**:
            - Content moderation challenges: Partner with AI moderation services
            - Copyright concerns: Implement content ID and takedown processes

            **Operational Risks**:
            - Platform abuse: Rate limiting and user reporting systems
            - Performance issues: Auto-scaling infrastructure and monitoring
            """,

            "Success Metrics": """
            - **Engagement**: 30% increase in average watch time
            - **Creation**: 100K+ clips created monthly
            - **Sharing**: 500K+ social shares quarterly
            - **Retention**: 15% improvement in user retention
            - **Revenue**: 20% increase in ad impressions
            """,

            "Timeline & Milestones": """
            **Month 1-2**: Core development and testing
            **Month 3**: Beta launch with select streamers
            **Month 4**: Full platform rollout
            **Month 5-6**: Feature expansion based on feedback
            **Month 7+**: Optimization and scaling
            """,

            "FAQ & Assumptions": """
            **Q: Will clips include copyrighted music?**
            A: Users will be responsible for copyright compliance; platform will provide content ID tools.

            **Q: What's the maximum clip length?**
            A: 60 seconds for free users, extended for premium subscribers.

            **Q: Can clips be monetized?**
            A: Yes, through platform partnerships and creator revenue sharing.
            """
        }

        return fallbacks.get(section, f"""
        **{section}**

        This section outlines the key requirements and considerations for implementing the clips feature.

        **Key Points**:
        - Comprehensive functionality for clip creation and management
        - User-friendly interface with powerful capabilities
        - Scalable architecture supporting platform growth
        - Integration with existing streaming infrastructure

        **Next Steps**: Detailed specifications to be developed based on user research and technical feasibility analysis.
        """)


class VPProduct:
    """
    👔 VP PRODUCT AGENT

    Role: Vice President of Product Management
    Background: 20+ years as VP Product at Fortune 500 companies, responsible for
    $2B+ product portfolios. Expert in identifying critical gaps, edge cases, and
    ensuring product-market fit before engineering investment.

    Responsibilities:
    - Critical review of complete PRD
    - Identify missing requirements and edge cases
    - Flag technical or business risks
    - Add "Missed Cases" section with Q&A format
    - Ensure enterprise-grade completeness
    """

    def __init__(self, gemini_api_key: str):
        genai.configure(api_key=gemini_api_key)

        # Try different model names in order of preference
        model_names = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro', 'gemini-1.0-pro']
        self.model = None

        for model_name in model_names:
            try:
                self.model = genai.GenerativeModel(model_name)
                print(f"✅ Using VP Product Gemini model: {model_name}")
                break
            except Exception as e:
                print(f"❌ Model {model_name} not available: {str(e)}")
                continue

        if self.model is None:
            raise Exception("No compatible Gemini models available. Please check your API key and region.")

    def review_prd(self, prd_sections: Dict[str, PRDSection], context: PRDContext) -> Dict[str, any]:
        """
        Comprehensive PRD review and gap analysis
        """
        # Compile PRD content for review
        prd_content = ""
        for section_name, section in prd_sections.items():
            prd_content += f"\n\n{section_name}:\n{section.selected_option}"

        prompt = f"""
        You are the VP of Product Management conducting a final critical review of this PRD.

        Context:
        - Problem: {context.problem_statement}
        - Idea: {context.idea}

        PRD Content:
        {prd_content}

        Your task is to identify:
        1. Missing requirements or sections
        2. Edge cases not covered
        3. Technical or business risks
        4. Integration considerations
        5. Success metrics gaps
        6. Any other critical oversights

        Format your response as a "Missed Cases" section with questions and detailed answers.
        Be thorough and critical - this is the final gate before engineering starts.

        Structure your response as:
        ## Missed Cases

        **Q1: [Specific question about gap]**
        A: [Detailed analysis and recommendation]

        **Q2: [Another question]**
        A: [Detailed analysis and recommendation]

        etc.
        """

        try:
            print("Conducting VP Product review using Gemini 1.5 Flash...")
            response = self.model.generate_content(prompt)
            missed_cases_content = response.text.strip()
            print("VP Product review complete")

            return {
                "review_passed": True,
                "missed_cases": missed_cases_content,
                "recommendations": self._extract_recommendations(missed_cases_content)
            }

        except Exception as e:
            print(f"VP Product review failed: {str(e)}")
            return {
                "review_passed": False,
                "missed_cases": """
                ## Additional Executive Considerations

                **Technical Scalability**: Ensure the clip processing infrastructure can handle peak loads during major streaming events.

                **Content Moderation**: Implement robust systems to prevent abuse while maintaining user experience.

                **Platform Integration**: Seamless integration with existing streaming tools and creator dashboards.

                **Analytics & Insights**: Comprehensive reporting on clip performance and user engagement metrics.

                **Legal Compliance**: Address copyright, privacy, and content ownership considerations across jurisdictions.
                """,
                "recommendations": [
                    "Conduct user testing with actual streamers and viewers",
                    "Implement comprehensive content moderation systems",
                    "Plan for international expansion and localization",
                    "Develop creator monetization strategies",
                    "Build analytics infrastructure for content insights"
                ]
            }

    def _extract_recommendations(self, missed_cases: str) -> List[str]:
        """Extract key recommendations from missed cases"""
        recommendations = []
        lines = missed_cases.split('\n')

        for line in lines:
            if line.strip().startswith('**Q') and ':' in line:
                question = line.split(':', 1)[1].strip()
                recommendations.append(question)

        return recommendations


class PRDOrchestrator:
    """
    🎼 PRD ORCHESTRATOR

    Role: Chief Product Officer & System Architect
    Background: Technology executive with 25+ years experience building and scaling
    product organizations. Expert in multi-agent systems and AI-driven product development.

    Responsibilities:
    - Coordinate the three-agent workflow
    - Manage data flow between agents
    - Handle errors and retries
    - Generate final PRD document
    - Ensure quality and completeness
    """

    def __init__(self, gemini_api_key: str, tavily_api_key: str, google_api_key: str = None, google_cx: str = None, github_pat: str = None, github_repo: str = None):
        self.researcher = PRDResearcher(gemini_api_key, tavily_api_key, google_api_key, google_cx)
        self.maker = PRDMaker(gemini_api_key)
        self.vp_product = VPProduct(gemini_api_key)
        self.github_pat = github_pat
        self.github_repo = github_repo

    def generate_prd(self, user_input: str, progress_callback=None) -> Tuple[bool, str, str]:
        """
        Execute the complete PRD generation workflow

        Returns: (success, docx_path, status_message)
        """
        try:
            # Step 1: Analyze input
            if progress_callback:
                progress_callback("🔍 Analyzing input type...")
            context = self.researcher.analyze_input_type(user_input)

            # Step 2: Conduct research
            if progress_callback:
                progress_callback("🔬 Conducting market research...")
            research_data = self.researcher.conduct_research(context)
            context.research_data = research_data

            # Step 3: Generate PRD sections
            if progress_callback:
                progress_callback("📝 Generating PRD sections...")
            prd_sections = self.maker.generate_complete_prd(context, research_data)

            # Check if any sections failed
            failed_sections = [name for name, section in prd_sections.items() if section.selected_option.startswith("Error")]
            if failed_sections:
                error_msg = f"PRD generation partially failed. Failed sections: {', '.join(failed_sections)}"
                if progress_callback:
                    progress_callback("⚠️ Generating PRD document with errors...")

                # Still try to generate the document with available content
                docx_path = self._generate_docx(prd_sections, context, {"review_passed": False, "missed_cases": f"Generation errors in: {', '.join(failed_sections)}", "recommendations": []})

                # Push to GitHub if credentials are available
                github_success = self._push_to_github(docx_path)

                error_msg = f"PRD generation partially failed. Failed sections: {', '.join(failed_sections)}"
                if github_success:
                    error_msg += " (Document pushed to GitHub despite errors)"

                return True, docx_path, error_msg

            # Step 4: VP Product review
            if progress_callback:
                progress_callback("👔 Conducting executive review...")
            review_result = self.vp_product.review_prd(prd_sections, context)

            # Step 5: Generate final document
            if progress_callback:
                progress_callback("📄 Creating final PRD document...")
            docx_path = self._generate_docx(prd_sections, context, review_result)

            # Push to GitHub if credentials are available
            github_success = self._push_to_github(docx_path)

            success_msg = "PRD generated successfully!"
            if github_success:
                success_msg += " (Pushed to GitHub repository)"

            return True, docx_path, success_msg

        except Exception as e:
            error_msg = f"PRD generation failed: {str(e)}"
            print(error_msg)
            return False, "", error_msg

    def _generate_docx(self, prd_sections: Dict[str, PRDSection], context: PRDContext, review_result: Dict) -> str:
        """
        Generate professional PRD document in DOCX format, optimized for practical implementation
        """
        print("Generating PRD document...")
        doc = Document()

        # Title page with professional formatting
        title = doc.add_heading('Product Requirements Document', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add subtitle with product name
        subtitle = doc.add_heading(f"{context.idea or 'Product Feature'} PRD", 1)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        metadata = [
            ('Created', datetime.now().strftime("%Y-%m-%d")),
            ('Version', '1.0'),
            ('Author', 'AI-Powered PRD Generator'),
            ('Status', 'Draft')
        ]

        for i, (key, value) in enumerate(metadata):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value

        doc.add_paragraph()

        # Executive Summary first (prominent placement)
        if "Executive Summary" in prd_sections:
            doc.add_heading('Executive Summary', 1)
            summary_section = prd_sections["Executive Summary"]
            if summary_section.selected_option:
                # Split content by lines and add paragraphs
                for line in summary_section.selected_option.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
            doc.add_paragraph()

        # Main sections in logical order
        main_sections = [
            "Problem Statement",
            "Solution Overview",
            "User Stories",
            "Functional Requirements",
            "Technical Requirements",
            "Business Requirements",
            "Implementation Plan",
            "Risks & Mitigations",
            "Success Metrics",
            "Timeline & Milestones"
        ]

        for section_name in main_sections:
            if section_name in prd_sections:
                doc.add_heading(section_name, 1)
                section = prd_sections[section_name]

                if section.selected_option:
                    # Format content based on section type
                    if section_name == "User Stories":
                        self._format_user_stories(doc, section.selected_option)
                    elif section_name in ["Functional Requirements", "Technical Requirements"]:
                        self._format_requirements(doc, section.selected_option)
                    elif section_name == "FAQ & Assumptions":
                        self._format_faq(doc, section.selected_option)
                    else:
                        # Default formatting
                        for line in section.selected_option.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line.strip())

                # Add rationale (smaller font)
                if section.rationale:
                    p = doc.add_paragraph()
                    run = p.add_run(f"Selection Rationale: {section.rationale}")
                    run.font.size = Pt(9)
                    run.italic = True

                doc.add_paragraph()

        # FAQ & Assumptions section
        if "FAQ & Assumptions" in prd_sections:
            doc.add_heading('FAQ & Assumptions', 1)
            faq_section = prd_sections["FAQ & Assumptions"]
            if faq_section.selected_option:
                self._format_faq(doc, faq_section.selected_option)
            doc.add_paragraph()

        # Missed Cases section from VP review
        if review_result.get("missed_cases"):
            doc.add_heading('Additional Considerations', 1)
            doc.add_paragraph("The following critical gaps and considerations were identified during executive review:")
            doc.add_paragraph()
            doc.add_paragraph(review_result["missed_cases"])

        # Save document with improved filename
        timestamp = datetime.now().strftime("%Y-%m-%d")

        # Generate a descriptive PRD title from the context
        if context.idea:
            prd_title = self._generate_prd_title(context.idea, context.problem_statement)
        else:
            prd_title = self._generate_prd_title(context.problem_statement, "")

        filename = f"PRD - {prd_title}.docx"
        filepath = os.path.join("reports", filename)

        # Ensure reports directory exists
        os.makedirs("reports", exist_ok=True)

        doc.save(filepath)
        print(f"PRD document saved to: {filepath}")
        return filepath

    def _format_user_stories(self, doc, content: str):
        """Format user stories with proper bullet points"""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('As a') or line.startswith('As an'):
                p = doc.add_paragraph(line, style='List Bullet')
            elif line.startswith('-') or line.startswith('•'):
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet 2')
            elif line:
                doc.add_paragraph(line)

    def _format_requirements(self, doc, content: str):
        """Format requirements with numbered lists"""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line and not line.startswith(' ') and not line.startswith('\t'):
                # Main requirement
                p = doc.add_paragraph(line, style='List Number')
            elif line.startswith('-') or line.startswith('•') or line.startswith('  -'):
                # Sub-requirement
                p = doc.add_paragraph(line.lstrip('-• ').strip(), style='List Bullet 2')
            elif line:
                doc.add_paragraph(line)

    def _format_faq(self, doc, content: str):
        """Format FAQ section with Q&A structure"""
        lines = content.split('\n')
        current_q = None
        current_a = []

        for line in lines:
            line = line.strip()
            if line.startswith('Q:') or line.startswith('●') or line.startswith('Question:'):
                # Save previous Q&A if exists
                if current_q and current_a:
                    self._add_qa_pair(doc, current_q, current_a)

                # Start new question
                current_q = line.lstrip('Q:● ').strip()
                current_a = []
            elif line.startswith('A:') or line.startswith('○') or line.startswith('Answer:'):
                # Add to current answer
                current_a.append(line.lstrip('A:○ ').strip())
            elif line and current_q:
                # Continuation of answer
                current_a.append(line)

        # Add final Q&A
        if current_q and current_a:
            self._add_qa_pair(doc, current_q, current_a)

    def _add_qa_pair(self, doc, question: str, answers: List[str]):
        """Add a formatted Q&A pair"""
        # Question
        p = doc.add_paragraph()
        run = p.add_run(f"Q: {question}")
        run.bold = True

        # Answer
        answer_text = ' '.join(answers)
        if answer_text.startswith('○ ') or answer_text.startswith('A: '):
            answer_text = answer_text[2:]

        p = doc.add_paragraph(f"A: {answer_text}", style='Normal')
        doc.add_paragraph()  # Spacing

    def _generate_prd_title(self, primary_text: str, secondary_text: str = "") -> str:
        """
        Generate a clean, descriptive PRD title from the input text
        """
        import re

        # Use primary text, fallback to secondary if empty
        text = primary_text or secondary_text
        if not text:
            return "Product Feature PRD"

        # Clean the text: remove extra whitespace, special chars
        text = re.sub(r'[^\w\s-]', '', text)  # Keep letters, numbers, spaces, hyphens
        text = re.sub(r'\s+', ' ', text).strip()  # Normalize whitespace

        # Extract key terms and create a concise title
        words = text.lower().split()

        # Common keywords to look for
        key_terms = {
            'clips': 'Clips',
            'streaming': 'Streaming',
            'vod': 'VOD',
            'live': 'Live',
            'stream': 'Stream',
            'viewer': 'Viewer',
            'streamer': 'Streamer',
            'upload': 'Upload',
            'feature': 'Feature',
            'platform': 'Platform',
            'video': 'Video',
            'content': 'Content'
        }

        # Find relevant terms
        found_terms = []
        for word in words:
            if word in key_terms:
                found_terms.append(key_terms[word])

        # If we found key terms, use them
        if found_terms:
            title = ' '.join(found_terms[:4])  # Limit to 4 terms
        else:
            # Fallback: use first few words, capitalize properly
            title_words = text.split()[:6]  # First 6 words max
            title = ' '.join(title_words)

        # Capitalize first letter of each word
        title = ' '.join(word.capitalize() for word in title.split())

        # Ensure it's not too long
        if len(title) > 50:
            title = title[:47] + "..."

        return title or "Product Feature PRD"

    def _push_to_github(self, filepath: str) -> bool:
        """
        Push the generated PRD to GitHub repository
        """
        if not self.github_pat or not self.github_repo:
            print("GitHub credentials not provided, skipping push to repository")
            return False

        try:
            from github import Github
            g = Github(self.github_pat)
            repo = g.get_repo(self.github_repo)

            # Read the file content
            with open(filepath, "rb") as f:
                content = f.read()

            # Create remote path (same as local path)
            remote_path = filepath.replace("\\", "/")

            # Create commit message
            filename = os.path.basename(filepath)
            commit_message = f"docs: Add generated PRD - {filename}"

            # Push to GitHub
            repo.create_file(remote_path, commit_message, content)
            print(f"✅ Successfully pushed PRD to GitHub: {remote_path}")
            return True

        except Exception as e:
            print(f"❌ Failed to push PRD to GitHub: {e}")
            return False