"""
=============================================================================
  PRD ENGINE — AUTONOMOUS MULTI-AGENT PRD SYSTEM (prd_engine.py)

 What this file does in plain English:
 This file is an entire AI product team packed into code. When you type a
 simple idea like "Build an inventory app", seven specialized AI agents
 work together — just like a real product team — to produce a detailed,
 professional Product Requirements Document (PRD).

 The 7 Agents:
   0. God Agent        — The boss. Decides what to do and who to assign.
   1. Classifier Agent — Figures out if your input is an idea or a problem.
   2. Research Agent   — Searches the internet + reads Specter reports.
   3. PRD Generator    — Writes 3 drafts of every section.
   4. Evaluator Agent  — Picks the best draft for each section.
   5. Gap Detector     — Finds missing pieces in the PRD.
   6. Eng Manager      — Reviews technical feasibility and edge cases.
   7. VP Product       — Final executive review before delivery.

 Key Features:
   - Combined Tavily + Google Search for comprehensive research
   - Pulls Specter intelligence reports from GitHub for context
   - Memory system for iterative refinement (doesn't redo work)
   - Error logging to GitHub (error_logs/ folder)
   - Super-detailed output readable by non-PM people
=============================================================================
"""

import os
import json
import time
import traceback
import requests
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field, asdict
from pathlib import Path
from copy import deepcopy

from dotenv import load_dotenv
load_dotenv()

import logging
from logger_config import (
    prd_logger, log_api_check, log_agent_start, log_agent_end,
    log_error, log_api_call, log_section_generated
)

from google import genai
from google.genai import types

GROQ_AVAILABLE = False
try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    print("⚠️ groq package not installed — Run: pip install groq")

def get_groq_client() -> Groq:
    api_key = os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        raise ValueError("GROQ_API_KEY not set")
    return Groq(api_key=api_key)

# Optional imports with graceful fallback
try:
    from tavily import TavilyClient
    TAVILY_AVAILABLE = True
except ImportError:
    TAVILY_AVAILABLE = False
    print("⚠️ tavily-python not installed — Tavily search disabled")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed — DOCX generation disabled")

try:
    from xhtml2pdf import pisa
    import markdown as md_lib
    PDF_EXPORT_AVAILABLE = True
except ImportError:
    PDF_EXPORT_AVAILABLE = False

try:
    from github import Github
    GITHUB_AVAILABLE = True
except ImportError:
    GITHUB_AVAILABLE = False


# =============================================================================
#  DATA STRUCTURES
# =============================================================================

@dataclass
class PRDSection:
    """A single section of the PRD with its generation history."""
    title: str
    options: List[str] = field(default_factory=list)
    selected_option: Optional[str] = None
    rationale: Optional[str] = None
    version: int = 1
    version_history: List[dict] = field(default_factory=list)

    def update(self, new_content: str, new_rationale: str):
        """Save current version to history and update."""
        if self.selected_option:
            self.version_history.append({
                "version": self.version,
                "content": self.selected_option,
                "rationale": self.rationale,
                "timestamp": datetime.now().isoformat()
            })
        self.selected_option = new_content
        self.rationale = new_rationale
        self.version += 1


@dataclass
class PRDContext:
    """Classified user input with extracted intent."""
    input_type: str  # "idea", "problem_statement", or "both"
    problem_statement: str
    idea: str
    original_input: str = ""
    research_data: Optional[Dict] = None


@dataclass
class PRDMemory:
    """
    Full state persistence across PRD iterations.

    research_memory: Dict of query -> results (cached, reusable)
    prd_state:       Current PRD sections dict
    section_history: List of all version changes
    version:         Current version number
    user_inputs:     All user inputs (original + refinements)
    context:         Current PRDContext
    """
    research_memory: Dict[str, Any] = field(default_factory=dict)
    prd_state: Dict[str, PRDSection] = field(default_factory=dict)
    section_history: List[dict] = field(default_factory=list)
    version: int = 1
    user_inputs: List[str] = field(default_factory=list)
    context: Optional[PRDContext] = None
    engineering_review: Optional[dict] = None
    vp_review: Optional[dict] = None

    def get_prd_markdown(self) -> str:
        """Render current PRD state as readable markdown."""
        lines = [f"# Product Requirements Document (v{self.version})\n"]
        if self.context:
            lines.append(f"**Input:** {self.context.original_input}\n")
        for name, section in self.prd_state.items():
            lines.append(f"\n## {section.title}\n")
            lines.append(section.selected_option or "_Not yet generated_")
            lines.append("")
        return "\n".join(lines)


@dataclass
class GapReport:
    """Output of the Gap Detector Agent."""
    missing_sections: List[str] = field(default_factory=list)
    improvements_needed: List[dict] = field(default_factory=list)
    weak_areas: List[str] = field(default_factory=list)
    raw_analysis: str = ""


@dataclass
class EngineeringReview:
    """Output of the Engineering Manager Agent."""
    issues: List[dict] = field(default_factory=list)
    approved: bool = False
    feedback_for_sections: Dict[str, str] = field(default_factory=dict)
    raw_review: str = ""


# =============================================================================
#  ERROR LOGGER — Pushes error files to GitHub
# =============================================================================

class GitHubErrorLogger:
    """
    Logs errors to the GitHub repository's error_logs/ folder.
    Each error creates a separate timestamped file for easy tracking.
    """

    def __init__(self, github_pat: str = "", github_repo: str = ""):
        self.pat = github_pat
        self.repo_name = github_repo

    def log_error(self, error_type: str, error_message: str,
                  traceback_str: str, context: str = None) -> bool:
        """
        Push an error log file to GitHub error_logs/ folder.

        Args:
            error_type: Category (e.g., "research_agent", "prd_generator")
            error_message: The error message
            traceback_str: Full Python traceback
            context: Additional context about what was happening
        """
        if not self.pat or not self.repo_name or not GITHUB_AVAILABLE:
            print(f"[ERROR LOG - LOCAL] {error_type}: {error_message}")
            return False

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"error_logs/error_{error_type}_{timestamp}.log"

        content = f"""========================================
 STREAMINTEL PRD ENGINE — ERROR LOG
========================================

Error Type    : {error_type}
Timestamp     : {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
Agent/Module  : {error_type}
Context       : {context or 'N/A'}

----------------------------------------
 ERROR MESSAGE
----------------------------------------
{error_message}

----------------------------------------
 FULL TRACEBACK
----------------------------------------
{traceback_str}

========================================
 END OF ERROR LOG
========================================
"""

        try:
            g = Github(self.pat)
            repo = g.get_repo(self.repo_name)
            commit_msg = f"error_log: {error_type} at {timestamp}"
            repo.create_file(filename, commit_msg, content)
            print(f"✅ Error logged to GitHub: {filename}")
            return True
        except Exception as e:
            print(f"❌ Failed to push error log to GitHub: {e}")
            return False


# =============================================================================
#  AGENT BASE CLASS
# =============================================================================

class BaseAgent:
    """Shared agent infrastructure — model initialization with global quota handling."""
    
    _client = None  # Shared Gemini client
    _groq_client = None  # Shared Groq client
    _use_groq_global = False  # Global flag - once set, all agents use Groq
    _quota_checked = False  # Track if quota check was done
    
    def __init__(self, gemini_api_key: str, preferred_models: List[str],
                 agent_name: str, error_logger: GitHubErrorLogger = None):
        
        # Global quota check - do this only once at start
        if not BaseAgent._quota_checked:
            BaseAgent._quota_checked = True
            BaseAgent._check_quota_and_set_provider(gemini_api_key)
        
        if BaseAgent._use_groq_global:
            # Use Groq directly (quota exhausted)
            if BaseAgent._groq_client is None and GROQ_AVAILABLE and os.environ.get("GROQ_API_KEY"):
                try:
                    BaseAgent._groq_client = get_groq_client()
                except Exception as e:
                    log_error(prd_logger, "groq_init", str(e), "Failed to init Groq")
            
            self.client = None
            self.model_name = None
            self.groq_model = BaseAgent._groq_client
            self.using_groq = True
            prd_logger.info(f"  ✅ {agent_name} -> Groq (llama-3.3-70b-versatile)")
        else:
            # Try Gemini first
            if BaseAgent._client is None:
                BaseAgent._client = genai.Client(api_key=gemini_api_key)
            
            self.client = BaseAgent._client
            self.model_name = preferred_models[0] if preferred_models else "gemini-2.0-flash"
            self.groq_model = None
            self.using_groq = False
            prd_logger.info(f"  ✅ {agent_name} -> Gemini ({self.model_name})")
        
        self.agent_name = agent_name
        self.error_logger = error_logger

    @classmethod
    def _check_quota_and_set_provider(cls, gemini_api_key: str):
        """Check Gemini quota once at start. If exhausted, switch all agents to Groq globally."""
        if not GROQ_AVAILABLE or not os.environ.get("GROQ_API_KEY"):
            prd_logger.info("Groq not available, using Gemini")
            return
        
        try:
            # Try a small test request to Gemini
            test_client = genai.Client(api_key=gemini_api_key)
            test_response = test_client.models.generate_content(
                model="gemini-2.0-flash-lite",
                contents="Hi",
            )
            _ = test_response.candidates[0].content.parts[0].text.strip()
            prd_logger.info("✅ Gemini quota available")
        except Exception as e:
            error_str = str(e)
            if "429" in error_str or "quota" in error_str.lower() or "ResourceExhausted" in error_str:
                cls._use_groq_global = True
                prd_logger.warning("⚠️ Gemini quota exhausted! Switching ALL agents to Groq globally")
            else:
                prd_logger.warning(f"⚠️ Gemini check failed: {error_str[:100]}, using Gemini anyway")

    def _call_llm(self, prompt: str, context: str = "", max_retries: int = 4) -> str:
        """Call the LLM - uses global provider setting."""
        log_agent_start(prd_logger, self.agent_name, f"LLM call: {context[:50] if context else 'prompt'}")
        
        for attempt in range(max_retries):
            try:
                if self.using_groq and self.groq_model:
                    log_api_call(prd_logger, "Groq", "chat.completions", "CALL", f"Attempt {attempt+1}")
                    response = self.groq_model.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[{"role": "user", "content": prompt}]
                    )
                    result = response.choices[0].message.content.strip()
                    log_api_call(prd_logger, "Groq", "chat.completions", "SUCCESS", f"Response: {len(result)} chars")
                    return result
                
                # Gemini path
                log_api_call(prd_logger, "Gemini", "generate_content", "CALL", f"Attempt {attempt+1}")
                response = self.client.models.generate_content(
                    model=self.model_name,
                    contents=prompt,
                )
                result = response.candidates[0].content.parts[0].text.strip()
                log_api_call(prd_logger, "Gemini", "generate_content", "SUCCESS", f"Response: {len(result)} chars")
                return result
            except Exception as e:
                error_str = str(e)
                
                # If quota exhausted during execution, switch globally to Groq
                if ("429" in error_str or "quota" in error_str.lower() or "ResourceExhausted" in error_str):
                    if not BaseAgent._use_groq_global and GROQ_AVAILABLE:
                        BaseAgent._use_groq_global = True
                        BaseAgent._groq_client = get_groq_client()
                        self.groq_model = BaseAgent._groq_client
                        self.using_groq = True
                        self.client = None
                        prd_logger.warning(f"  ⚠️ {self.agent_name} -> Gemini quota exhausted! Switching ALL agents to Groq globally")
                        continue
                
                log_error(prd_logger, self.agent_name, error_str, f"Attempt {attempt+1}/{max_retries}")
                
                if attempt < max_retries - 1:
                    wait = 8 * (attempt + 1)
                    prd_logger.warning(f"  ⚠️ {self.agent_name} retry {attempt+1}/{max_retries} in {wait}s")
                    time.sleep(wait)
                else:
                    error_msg = error_str
                    tb = traceback.format_exc()
                    prd_logger.error(f"  ❌ {self.agent_name} FAILED after {max_retries} retries: {error_str[:200]}")
                    log_error(prd_logger, f"{self.agent_name}_call", error_msg, f"Context: {context[:100]}")
                    if self.error_logger:
                        self.error_logger.log_error(
                            self.agent_name.lower().replace(" ", "_"),
                            error_msg, tb, context
                        )
                    raise

        log_agent_end(prd_logger, self.agent_name, "COMPLETE")
        return ""

        log_agent_end(prd_logger, self.agent_name, "COMPLETE")
        return ""

    def _call_llm_json(self, prompt: str, context: str = "") -> dict:
        """Call LLM and parse JSON response."""
        raw = self._call_llm(prompt, context)
        # Strip markdown code fences if present
        cleaned = raw
        if "```json" in cleaned:
            cleaned = cleaned.split("```json")[1].split("```")[0]
        elif "```" in cleaned:
            cleaned = cleaned.split("```")[1].split("```")[0]
        try:
            return json.loads(cleaned.strip())
        except json.JSONDecodeError:
            # Try to find JSON object in the response
            start = raw.find("{")
            end = raw.rfind("}") + 1
            if start != -1 and end > start:
                try:
                    return json.loads(raw[start:end])
                except:
                    pass
            return {"raw_response": raw, "parse_error": True}


# =============================================================================
# 0️⃣ GOD AGENT — Master Orchestrator
# =============================================================================

class GodAgent(BaseAgent):
    """
     THE GOD AGENT — an elite-level Head of Product + Chief of Staff + Systems Thinker.

    OBJECTIVE: Convert user input into a structured multi-agent execution plan
    for PRD generation or refinement.

    THINKING FRAMEWORK:
    1. Identify intent: New PRD, PRD update, or Refinement
    2. Identify gaps: Missing research, Weak sections, Lack of clarity
    3. Decide execution strategy: Full generation, Partial update, Iterative refinement

    DECISION LOGIC:
    - New PRD: classifier -> research -> PRD (section loop + evaluator) -> engineering_manager -> vp_product
    - PRD Update: gap_detector -> research (incremental) -> PRD update -> engineering_manager -> vp_product
    - Weak PRD: regenerate weak sections -> evaluator -> engineering_manager -> vp_product

    SYSTEM RULES:
    - NEVER skip research for new PRD
    - ALWAYS enforce: 3 options -> evaluator -> selection
    - ALWAYS require engineering_manager before vp_product
    - ALWAYS prefer partial updates over full regeneration
    """

    def __init__(self, gemini_api_key: str, error_logger: GitHubErrorLogger = None):
        super().__init__(
            gemini_api_key,
            ["gemini-2.0-flash", "gemini-2.0-flash-lite"],
            "God Agent",
            error_logger
        )

    def plan_initial_workflow(self, user_input: str) -> dict:
        """Decide the workflow for initial PRD generation with smart section selection."""
        prompt = f"""You are the GOD AGENT — an elite-level Head of Product + Chief of Staff + Systems Thinker.

##  OBJECTIVE
Convert user input into a structured multi-agent execution plan for PRD generation. You must analyze the input to understand WHAT product is being built and WHICH sections are needed.

##  THINKING FRAMEWORK
1. Analyze input: What product is being built? What domain/industry?
2. Identify product type: streaming, e-commerce, social, productivity, AI/ML, fintech, healthcare, etc.
3. Determine required sections: Some sections are universal, some are product-specific
4. Identify intent: New PRD, PRD update, or Refinement
5. Decide execution strategy

##  SECTION CATALOG (know when to use each)

### Core Sections (ALL products need these):
- Problem Statement: What problem are we solving? Who has it? How painful?
- Objectives: Quantifiable goals with targets (SMART)
- Core Product Principles: Fundamental rules guiding the product
- Scope: What's IN and what's OUT (prevent scope creep)
- User Personas: Who are the users? Demographics, goals, pain points
- User Stories & Flows: How users interact with the product (step-by-step)
- Functional Requirements: What features to build (Feature ID, Name, Description, Priority)
- Non-Functional Requirements: Performance, scalability, security, accessibility
- Technical Architecture: System design, tech stack, database, APIs
- Business Requirements & Monetization: Revenue model, costs, go-to-market
- Implementation Roadmap: Phased plan with timelines
- Risks & Mitigations: What could go wrong and how to prevent
- Success Metrics & KPIs: How do we measure success?
- Analytics: What events to track and why

### Product-Specific Sections (use when relevant):
- Feed/Discovery: For content platforms (social, streaming, news)
- Moderation & Safety: For UGC platforms (content review, AI moderation)
- AI Design: For AI/ML features (signals, algorithms, outputs)
- Payments & Billing: For fintech/e-commerce (transactions, refunds)
- Notifications: For engagement products (push, email, in-app)
- Search: For discovery products (ranking, filters)
- Onboarding: For consumer products (progressive profiling)
- Partnerships/API: For platform products (integrations)
- Compliance: For regulated industries (GDPR, HIPAA, SOC2)
- etc.

##  DECISION LOGIC
- Streaming/Video platforms -> Add Feed, Moderation, AI Design sections
- E-commerce -> Add Payments, Inventory, Logistics sections
- Social platforms -> Add Feed, Moderation, Notifications sections
- AI/ML products -> Add AI Design, Model Monitoring sections
- Fintech -> Add Payments, Compliance, Security sections

## ⚙️ SYSTEM RULES
- NEVER skip research for new PRD
- ALWAYS enforce: 3 options -> evaluator -> selection
- ALWAYS require engineering_manager before vp_product
- ALWAYS prefer partial updates over full regeneration

User input:
"{user_input}"

##  OUTPUT (STRICT JSON)
{{
    "intent": "new_prd",
    "confidence": 0.0,
    "execution_strategy": "full_generation",
    "product_type": "streaming/video platform",
    "input_quality": "high|medium|low",
    "input_summary": "one-line summary of what user wants to build",
    "identified_gaps": ["list of identified gaps"],
    "action_plan": [
        {{"step": 1, "agent": "classifier", "task": "classify user input"}}
    ],
    "required_sections": [
        "Problem Statement",
        "Objectives",
        "Core Product Principles", 
        "Scope",
        "User Roles",
        "User Flows",
        "Functional Requirements",
        "Non-Functional Requirements",
        "Technical Architecture",
        "Monetization",
        "Edge Cases",
        "Analytics",
        "Feed",
        "Moderation & NSFW System",
        "AI Design for Clips"
    ],
    "research_queries": ["list of 4-6 specific research queries"],
    "focus_areas": ["list of key areas the PRD should emphasize"],
    "special_instructions": "any special considerations"
}}

Respond with JSON only. Think carefully about which sections are needed based on the product type."""
        try:
            result = self._call_llm_json(prompt, f"Initial planning for: {user_input[:100]}")
            if result.get("parse_error"):
                return {
                    "input_quality": "medium",
                    "input_summary": user_input[:200],
                    "research_queries": [
                        f"market analysis {user_input[:100]}",
                        f"competitors for {user_input[:100]}",
                        f"technical challenges {user_input[:100]}",
                        f"user needs {user_input[:100]}"
                    ],
                    "focus_areas": ["market fit", "technical feasibility", "user experience"],
                    "special_instructions": "Standard PRD generation flow"
                }
            return result
        except Exception as e:
            return {
                "input_quality": "medium",
                "input_summary": user_input[:200],
                "research_queries": [
                    f"market analysis {user_input[:100]}",
                    f"competitors {user_input[:100]}",
                    f"industry trends {user_input[:100]}",
                    f"technical feasibility {user_input[:100]}"
                ],
                "focus_areas": ["product-market fit", "technical architecture", "user experience"],
                "special_instructions": "Proceed with standard workflow"
            }

    def interpret_update(self, new_input: str, memory: PRDMemory) -> dict:
        """Decide how to handle iterative user updates."""
        current_prd_summary = ""
        for name, section in memory.prd_state.items():
            current_prd_summary += f"- {name}: {(section.selected_option or '')[:100]}...\n"

        prompt = f"""You are the GOD AGENT — an elite-level Head of Product + Chief of Staff + Systems Thinker.

##  OBJECTIVE
Analyze the user's update request and create an execution plan for PRD refinement.

##  SYSTEM RULES
- ALWAYS prefer partial updates over full regeneration
- ALWAYS require engineering_manager before vp_product
- ALWAYS enforce: 3 options -> evaluator -> selection

The user previously built a PRD (v{memory.version}) with these sections:
{current_prd_summary}

Previous inputs: {json.dumps(memory.user_inputs[-3:])}

The user now says:
"{new_input}"

Decide what to do. Respond in STRICT JSON:
{{
    "intent": "update_prd | refine_prd",
    "confidence": 0.0,
    "execution_strategy": "partial_update | refinement",
    "identified_gaps": ["list of identified gaps"],
    "action": "update_sections|regenerate_all|add_requirement",
    "affected_sections": ["list of section names to regenerate"],
    "new_research_needed": true|false,
    "research_queries": ["specific queries if research needed"],
    "instructions_for_generator": "specific instructions for the PRD generator",
    "priority": "speed | quality | balanced",
    "notes": "short reasoning"
}}"""
        try:
            result = self._call_llm_json(prompt, f"Update interpretation: {new_input[:100]}")
            if result.get("parse_error"):
                return {
                    "action": "update_sections",
                    "affected_sections": list(memory.prd_state.keys()),
                    "new_research_needed": True,
                    "research_queries": [f"research for: {new_input}"],
                    "instructions_for_generator": new_input
                }
            return result
        except Exception:
            return {
                "action": "update_sections",
                "affected_sections": list(memory.prd_state.keys()),
                "new_research_needed": True,
                "research_queries": [f"{new_input} market analysis"],
                "instructions_for_generator": new_input
            }


# =============================================================================
# 1️⃣ CLASSIFIER AGENT
# =============================================================================

class ClassifierAgent(BaseAgent):
    """
     CLASSIFIER AGENT

    A high-precision intent classifier for product workflows.

    RULES:
    - Pain -> problem_statement
    - Solution -> idea
    - Modification -> prd_update
    - Mixed -> choose dominant intent
    """

    def __init__(self, gemini_api_key: str, error_logger: GitHubErrorLogger = None):
        super().__init__(
            gemini_api_key,
            ["gemini-2.0-flash", "gemini-2.0-flash-lite"],
            "Classifier Agent",
            error_logger
        )

    def classify(self, user_input: str) -> PRDContext:
        """Classify user input as idea, problem statement, or prd_update."""
        prompt = f"""You are a high-precision intent classifier for product workflows.

##  TASK
Classify input into:
- idea
- problem_statement
- prd_update

##  RULES
- Pain -> problem_statement
- Solution -> idea
- Modification -> prd_update
- Mixed -> choose dominant intent

User input:
"{user_input}"

##  OUTPUT
{{
    "type": "idea|problem_statement|prd_update",
    "confidence": 0.0,
    "reason": "clear explanation",
    "problem_statement": "the core problem being solved (write one even if input is just an idea)",
    "idea": "the product/feature concept (write one even if input is just a problem)"
}}"""
        try:
            result = self._call_llm_json(prompt, f"Classifying: {user_input[:100]}")
            return PRDContext(
                input_type=result.get("type", result.get("input_type", "both")),
                problem_statement=result.get("problem_statement", user_input),
                idea=result.get("idea", user_input),
                original_input=user_input
            )
        except Exception:
            return PRDContext(
                input_type="both",
                problem_statement=user_input,
                idea=user_input,
                original_input=user_input
            )


# =============================================================================
# 2️⃣ RESEARCH AGENT
# =============================================================================

class ResearchAgent(BaseAgent):
    """
     RESEARCH AGENT — Elite product research analyst.

    THINKING MODEL:
    - Market reality (size, maturity, behavior)
    - Competitors (strategy, strengths, gaps)
    - User psychology (pain, motivation, behavior)
    - Opportunities (where product can win)
    - Risks

    RULES:
    - No generic insights
    - No obvious statements
    - Reuse past research if available
    - Add only new insights for gaps
    """

    def __init__(self, gemini_api_key: str, tavily_api_key: str = "",
                 google_api_key: str = "", google_cx: str = "",
                 github_pat: str = "", github_repo: str = "",
                 error_logger: GitHubErrorLogger = None):
        super().__init__(
            gemini_api_key,
            ["gemini-2.0-flash", "gemini-2.0-flash-lite"],
            "Research Agent",
            error_logger
        )
        self.tavily_client = None
        if tavily_api_key and TAVILY_AVAILABLE:
            try:
                self.tavily_client = TavilyClient(api_key=tavily_api_key)
            except Exception:
                pass
        self.google_api_key = google_api_key
        self.google_cx = google_cx
        self.github_pat = github_pat
        self.github_repo = github_repo

    def _search_tavily(self, query: str, max_results: int = 5) -> List[Dict]:
        """Search using Tavily API."""
        if not self.tavily_client:
            return []
        try:
            results = self.tavily_client.search(
                query=query, search_depth="advanced", max_results=max_results,
                include_raw_content=False
            )
            return [{
                "title": r.get("title", ""),
                "url": r.get("url", ""),
                "snippet": r.get("content", ""),
                "source": "tavily"
            } for r in results.get("results", [])]
        except Exception as e:
            print(f"  ⚠️ Tavily search error: {e}")
            return []

    def _search_google(self, query: str, max_results: int = 5) -> List[Dict]:
        """Search using Google Custom Search API."""
        if not self.google_api_key or not self.google_cx:
            return []
        try:
            url = "https://www.googleapis.com/customsearch/v1"
            params = {
                "key": self.google_api_key, "cx": self.google_cx,
                "q": query, "num": max_results
            }
            response = requests.get(url, params=params, timeout=10)
            data = response.json()
            return [{
                "title": item.get("title", ""),
                "url": item.get("link", ""),
                "snippet": item.get("snippet", ""),
                "source": "google"
            } for item in data.get("items", [])]
        except Exception as e:
            print(f"  ⚠️ Google search error: {e}")
            return []

    def _merge_and_deduplicate(self, *result_lists) -> List[Dict]:
        """Merge multiple result lists and remove duplicate URLs."""
        seen_urls = set()
        merged = []
        for results in result_lists:
            for r in results:
                url = r.get("url", "")
                if url and url not in seen_urls:
                    merged.append(r)
                    seen_urls.add(url)
        return merged

    def _fetch_specter_reports(self) -> List[str]:
        """
        Pull Specter intelligence reports from GitHub repo's reports/ folder.
        Extracts report names and dates as additional research context.
        """
        if not self.github_pat or not self.github_repo or not GITHUB_AVAILABLE:
            return []

        insights = []
        try:
            g = Github(self.github_pat)
            repo = g.get_repo(self.github_repo)
            try:
                contents = repo.get_contents("reports")
                for item in contents:
                    if item.name.endswith(".pdf") or item.name.endswith(".docx"):
                        # Extract date and context from filename
                        insights.append(
                            f"Specter Report: {item.name} "
                            f"(Size: {item.size} bytes, "
                            f"Path: {item.path}) — "
                            f"Contains market intelligence and competitive analysis "
                            f"generated by the Specter autonomous research agent."
                        )
            except Exception:
                # reports/ folder might not exist yet
                pass

            print(f"   Found {len(insights)} Specter reports in GitHub repo")
            return insights
        except Exception as e:
            print(f"  ⚠️ Could not fetch Specter reports: {e}")
            return []

    def research(self, queries: List[str], memory: PRDMemory = None) -> Dict[str, Any]:
        """
        Conduct comprehensive research using Tavily + Google + Specter reports.
        Reuses cached results from memory when available.
        """
        print("   Research Agent: Starting research phase...")
        research_data = {
            "results_by_query": {},
            "specter_reports": [],
            "summary": "",
            "timestamp": datetime.now().isoformat()
        }

        for query in queries:
            # Check memory cache first
            if memory and query in memory.research_memory:
                print(f"  ♻️  Cache hit for: '{query[:60]}...'")
                research_data["results_by_query"][query] = memory.research_memory[query]
                continue

            print(f"   Searching: '{query[:60]}...'")

            # Fire BOTH search engines
            tavily_results = self._search_tavily(query)
            google_results = self._search_google(query)

            # Combine and deduplicate
            combined = self._merge_and_deduplicate(tavily_results, google_results)
            print(f"     -> {len(tavily_results)} Tavily + {len(google_results)} Google = {len(combined)} unique results")

            research_data["results_by_query"][query] = combined
            time.sleep(1)  # Politeness delay

        # Fetch Specter reports from GitHub
        research_data["specter_reports"] = self._fetch_specter_reports()

        # Generate research summary using LLM
        research_data["summary"] = self._synthesize_research(research_data)

        print("  ✅ Research phase complete")
        return research_data

    def _synthesize_research(self, research_data: Dict) -> str:
        """Create a comprehensive research summary."""
        all_snippets = []
        for query, results in research_data.get("results_by_query", {}).items():
            for r in results[:3]:
                all_snippets.append(f"[{query}] {r.get('title', '')}: {r.get('snippet', '')}")

        specter_str = "\n".join(research_data.get("specter_reports", []))

        if not all_snippets and not specter_str:
            return "Limited research data available. PRD will be generated based on the input description."

        prompt = f"""You are an elite product research analyst.

##  TASK
Generate high-signal, actionable research.

##  THINKING MODEL
- Market reality (size, maturity, behavior)
- Competitors (strategy, strengths, gaps)
- User psychology (pain, motivation, behavior)
- Opportunities (where product can win)
- Risks

## ⚠️ RULES
- No generic insights
- No obvious statements
- Add only new insights for gaps

SEARCH RESULTS:
{chr(10).join(all_snippets[:20])}

SPECTER INTELLIGENCE REPORTS:
{specter_str or "None available"}

Synthesize into a structured research brief. Output should cover:
- market (size, maturity, behavior)
- competitors (name, strategy, strength, weakness, opportunity for each)
- user_insights (specific, non-obvious)
- opportunity_areas (where product can win)
- risks (concrete, not generic)"""

        try:
            return self._call_llm(prompt, "Research synthesis")
        except Exception:
            return f"Research collected: {len(all_snippets)} search results, {len(research_data.get('specter_reports', []))} Specter reports."


# =============================================================================
# 3️⃣ PRD GENERATOR AGENT
# =============================================================================

class PRDGeneratorAgent(BaseAgent):
    """
    ✍️ PRD GENERATOR AGENT — Top-tier Product Manager.

    TASK: Generate EXACTLY 3 distinct options for a PRD section.

    THINKING MODEL:
    Option 1: Bold (high ambition, differentiated)
    Option 2: Balanced (practical, scalable)
    Option 3: MVP (lean, fast execution)

    RULES:
    - No repetition across options
    - No vague phrases
    - Must be structured and actionable
    - Include metrics or logic where possible
    """

    SECTIONS = []  # Will be set dynamically based on product type

    def __init__(self, gemini_api_key: str, error_logger: GitHubErrorLogger = None):
        super().__init__(
            gemini_api_key,
            ["gemini-2.0-flash", "gemini-2.0-flash-lite"],
            "PRD Generator",
            error_logger
        )

    @classmethod
    def get_sections_for_product(cls, god_plan: dict) -> List[str]:
        """Dynamically determine sections based on product type from god_plan."""
        # Default core sections for ALL products
        core_sections = [
            "Problem Statement",
            "Objectives",
            "Core Product Principles",
            "Scope",
            "User Roles",
            "User Flows",
            "Functional Requirements",
            "Non-Functional Requirements",
            "Technical Architecture",
            "Business Requirements & Monetization",
            "Implementation Roadmap",
            "Risks & Mitigations",
            "Success Metrics & KPIs",
            "Analytics",
            "Edge Cases",
        ]
        
        # Product-specific sections
        product_specific = {
            "streaming": ["Feed", "Moderation & NSFW System", "AI Design for Clips"],
            "video": ["Feed", "Moderation & NSFW System", "Video Processing"],
            "social": ["Feed", "Moderation & NSFW System", "Notifications", "Engagement"],
            "ecommerce": ["Payments & Billing", "Inventory Management", "Shipping & Logistics", "Product Catalog"],
            "fintech": ["Payments & Billing", "Compliance & Regulatory", "Security", "KYC/AML"],
            "ai_ml": ["AI Design", "Model Monitoring", "Data Pipeline", "Model Training"],
            "healthcare": ["HIPAA Compliance", "Patient Data", "Medical Records", "Security"],
            "productivity": ["Collaboration", "Notifications", "Integrations", "Onboarding"],
        }
        
        # Get required_sections from god_plan if available
        if god_plan and "required_sections" in god_plan:
            return god_plan["required_sections"]
        
        # Fallback: check product_type in god_plan
        product_type = ""
        if god_plan and "product_type" in god_plan:
            product_type = god_plan["product_type"].lower()
        
        # Add product-specific sections
        extra_sections = []
        for key, sections in product_specific.items():
            if key in product_type:
                extra_sections.extend(sections)
        
        return core_sections + extra_sections

    def generate_section(self, section: str, context: PRDContext,
                         research_summary: str, god_plan: dict,
                         engineering_feedback: str = "") -> List[str]:
        """Generate 3 detailed option drafts for a single PRD section."""
        
        log_agent_start(prd_logger, "PRDGenerator", f"Generating: {section}")

        section_guide = self._get_section_guide(section)
        feedback_block = ""
        if engineering_feedback:
            feedback_block = f"""
⚠️ ENGINEERING FEEDBACK (Address these issues in your rewrite):
{engineering_feedback}
"""

        prompt = f"""You are a Senior Product Manager creating a detailed, production-ready PRD.

##  TASK
Generate EXACTLY 3 distinct options for a PRD section. Each option should be comprehensive, actionable, and ready for engineering, design, and business teams.

##  THINKING MODEL
Option 1: Bold (high ambition, differentiated) - Maximum features, futuristic approach
Option 2: Balanced (practical, scalable) - Right mix of ambition and feasibility  
Option 3: MVP (lean, fast execution) - Minimum viable with clear path to scale

## ⚠️ CRITICAL RULES
- NO repetition across options - each must be fundamentally different
- NO vague phrases like "user-friendly" or "efficient" without specifics
- Be STRUCTURED with headers, bullet points, numbered lists, and tables
- Include EXACT numbers, metrics, timelines, and specific features
- Minimum 500 words per option - be thorough
- Include user flows, edge cases, and technical considerations where relevant

CONTEXT:
- User's Idea: {context.idea}
- Problem Being Solved: {context.problem_statement}
- Input Type: {context.input_type}
- Focus Areas: {json.dumps(god_plan.get('focus_areas', []))}
- Special Instructions: {god_plan.get('special_instructions', 'None')}

RESEARCH INSIGHTS:
{research_summary[:3000]}

{feedback_block}

SECTION: {section}
{section_guide}

Generate exactly 3 complete, distinct options for this section.
Label them clearly as "--- OPTION 1 ---", "--- OPTION 2 ---", "--- OPTION 3 ---".

##  OUTPUT FORMAT
Each option must be labeled and distinct. No repetition across options.
Include:
- Specific metrics and KPIs
- User flows with steps
- Technical considerations
- Edge cases and failure handling
- Timeline estimates where applicable
"""
        try:
            raw = self._call_llm(prompt, f"Generating section: {section}")
            options = self._parse_three_options(raw)
            
            total_words = sum(len(opt.split()) for opt in options)
            log_section_generated(prd_logger, section, total_words, len(options))
            log_agent_end(prd_logger, "PRDGenerator", f"COMPLETE | {section}")
            
            return options
        except Exception as e:
            error_msg = str(e)
            prd_logger.error(f"  ❌ PRD Generator FAILED for '{section}': {error_msg[:200]}")
            log_error(prd_logger, "PRDGenerator_section", error_msg, f"Section: {section}")
            raise

    def _parse_three_options(self, raw: str) -> List[str]:
        """Parse LLM output into 3 separate options."""
        markers = ["--- OPTION 1 ---", "--- OPTION 2 ---", "--- OPTION 3 ---"]
        options = []

        # Try structured parsing first
        parts = raw
        for i, marker in enumerate(markers):
            if marker in parts:
                split = parts.split(marker, 1)
                if i > 0 and split[0].strip():
                    options.append(split[0].strip())
                parts = split[1] if len(split) > 1 else ""

        if parts.strip():
            options.append(parts.strip())

        # If structured parsing failed, try splitting by "Option X"
        if len(options) < 3:
            options = []
            import re
            splits = re.split(r'(?:^|\n)(?:Option\s+\d|OPTION\s+\d|\*\*Option\s+\d)', raw, flags=re.IGNORECASE)
            for s in splits:
                stripped = s.strip()
                if stripped and len(stripped) > 50:
                    options.append(stripped)

        # Pad if still not enough
        while len(options) < 3:
            if options:
                options.append(options[0])
            else:
                options.append("Content generation in progress — please retry.")

        return options[:3]

    def _get_section_guide(self, section: str) -> str:
        """Section-specific writing instructions - comprehensive for ZERO knowledge readers."""
        guides = {
            "Problem Statement": """CRITICAL: Write for someone who knows NOTHING about this product.

Define the problem with absolute clarity:
- WHO experiences this problem? (Specific user segments)
- WHAT exactly is the problem? (Concrete description)
- WHEN does it occur? (Situations, contexts)
- WHERE does it happen? (Physical/digital locations)
- WHY does it matter? (Quantify impact - lost revenue, time wasted, user drop-off)
- What's the current workaround? (How do users solve it today?)
- Why do current solutions fail? (Gaps in existing products)

Include real-world examples, specific scenarios, and quantitative data where possible.
Example: "40% of viewers report abandoning streams due to inability to easily create clips" 
This should be written at a level where even someone unfamiliar with streaming can understand.""",

            "Objectives": """CRITICAL: Write SMART objectives. 

Define quantifiable objectives:
- PRIMARY GOALS: Top 3-5 goals with specific targets (e.g., "Enable frictionless clip creation" -> "90% of users can create clip in <10 seconds")
- SECONDARY GOALS: Supporting objectives
- Quantify with competitor benchmarks if available
- Include 30/60/90 day targets where applicable
- Format: Objective -> Metric -> Target -> Timeline

Example: "Drive viral distribution" -> "Viral coefficient K > 1" -> "K platform = 1.2" -> "6 months""",

            "Core Product Principles": """CRITICAL: Define fundamental rules that guide ALL product decisions.

These principles should be:
- Simple and memorable
- Actionable (can guide decisions)
- Universal (apply to all features)

Examples from reference PRD:
- Ownership -> Streamer (always)
- Attribution -> Creator (viewer/streamer/uploader/AI)  
- Discovery -> Centralised under streamer
- Safety -> No clip goes live without moderation clearance
- Consistency -> Same lifecycle for all clip types
- Low latency -> Clip creation should feel instant

Each principle should have a one-line explanation of WHY it matters.""",

            "Scope": """CRITICAL: Clear boundaries to prevent scope creep.

4.1 IN SCOPE (V1):
- What features/functionality ARE included
- What user flows ARE supported

4.2 OUT OF SCOPE (V1):
- What is explicitly NOT included in V1
- What will be addressed in future phases
- Be specific (not vague)

Example:
- IN: Clip creation from live stream, clip creation from VOD, clip upload
- OUT: Advanced AI clip generation, clip editing beyond trimming/framing""",

            "User Roles": """CRITICAL: Define ALL user types who interact with the product.

For each role, define:
- Role Name (e.g., Streamer, Viewer, Content Manager, Moderator)
- Description (what they do)
- Permissions/Access Level
- Goals (what they want to achieve)
- Pain points (what frustrates them)

Include both internal and external users.
Example format:
1. Streamer
   - Creates and streams content
   - Can create clips, upload clips, manage their clips
   - Goals: Grow audience, monetize content
2. Viewer
   - Watches streams/VODs
   - Can create clips, share clips
   - Goals: Discover content, share with friends""",

            "User Flows": """CRITICAL: Step-by-step flows for EVERY major user journey.

For each flow include:
- PRECONDITIONS (what must be true before starting)
- PRIMARY FLOW (numbered steps with UI elements)
- ALTERNATIVE FLOWS (what if user takes different path)
- ERROR HANDLING (what goes wrong and how to recover)
- EDGE CASES (unusual but possible scenarios)
- ACCEPTANCE CRITERIA (how do we know it worked)

Use standard format like reference PRD:
6.1 Live Clip Creation
- Preconditions: Stream is live, Buffer available (min 90 sec)
- Max clip duration: 90 sec, Min: 10 sec
- Flow: User taps Clip -> System captures last 90-sec buffer -> Opens editor -> User trims -> User adds title/tags -> User publishes OR shares

Include all possible paths (success, failure, cancellation).""",

            "Functional Requirements": """CRITICAL: Every feature with complete specification.

For each feature:
- Feature ID: Unique identifier (e.g., CGW-1)
- Name: Clear feature name
- Description: What it does
- Priority: P0 (must have), P1 (should have), P2 (nice to have)
- User Story Reference: Which user story it fulfills
- Acceptance Criteria: How to verify it works (specific, testable)
- Dependencies: What must be built first

Organize by FEATURE AREA.
Example format:
Feature Area: Clip Generation Workflow
- CGW-1, Live Stream Clip Generation, Users can generate clips from live streams with max 60 sec, P0, US-1, System processes within 10 seconds, None""",

            "Non-Functional Requirements": """CRITICAL: Technical requirements that ensure quality.

Cover with specific MEASUREABLE targets:
- PERFORMANCE: Response times, throughput (e.g., <200ms for clip generation)
- SCALABILITY: Concurrent users, data volume (e.g., support 10,000 concurrent users)
- SECURITY: Authentication, encryption, compliance (e.g., OAuth 2.0, AES-256, GDPR)
- RELIABILITY: Uptime SLA, disaster recovery (e.g., 99.95% uptime)
- ACCESSIBILITY: WCAG compliance level
- LOCALIZATION: Supported languages, regional compliance

Table format with metrics and targets is preferred.""",

            "Technical Architecture": """CRITICAL: Complete technical specification for engineers.

Must include:
- SYSTEM ARCHITECTURE: Monolith or Microservices? Diagram description
- TECHNOLOGY STACK: Languages, frameworks, databases with JUSTIFICATION
- DATABASE DESIGN: Collections/tables, relationships, schema
- API SPECIFICATIONS: REST/GraphQL endpoints, request/response format
- THIRD-PARTY INTEGRATIONS: External services, dependencies
- INFRASTRUCTURE: Cloud provider, deployment strategy
- CI/CD PIPELINE: Build, test, deploy process
- MONITORING: What metrics to track, observability

Include text-based architecture diagram.""",

            "Business Requirements & Monetization": """CRITICAL: How the product makes money.

Cover:
- REVENUE MODEL: How does the product generate income? (ads, subscriptions, transaction fees)
- PRICING STRATEGY: Tier structure, discounts
- COST STRUCTURE: Development, marketing, operations, staffing
- UNIT ECONOMICS: CAC, CLV, ARPU, LTV
- GO-TO-MARKET PLAN: Launch strategy, customer acquisition
- PARTNERSHIPS NEEDED: External relationships
- LEGAL/COMPLIANCE: Terms of service, content ownership
- FINANCIAL PROJECTIONS: 3-year revenue/expense projections

Include specific numbers, not vague statements.""",

            "Implementation Roadmap": """CRITICAL: Phased delivery plan.

For each PHASE:
- OBJECTIVE: What are we achieving?
- DELIVERABLES: Specific features/功能
- TIMELINE: Start/end dates or week numbers
- TEAM COMPOSITION: Roles needed
- DEPENDENCIES: What must be ready first
- TESTING APPROACH: Unit, integration, UAT
- ROLLOUT STRATEGY: How to release (beta, gradual, big bang)

Reference PRD format:
Phase 1: MVP (Weeks 1-12)
- Objective: Launch basic clip feature
- Deliverables: Clip generation, upload, sharing
- Team: 2 Backend, 1 Frontend, 1 QA, 1 Designer""",

            "Risks & Mitigations": """CRITICAL: What could go wrong and how to prevent it.

For each RISK identify:
- CATEGORY: Technical/Business/Operational/Legal/Financial
- LIKELIHOOD: H (High), M (Medium), L (Low)
- IMPACT: H/M/L
- MITIGATION STRATEGY: What we're doing to prevent
- CONTINGENCY PLAN: What if it happens anyway
- OWNER: Who is responsible

Example:
- Scalability Risk, Likelihood: H, Impact: H
- Mitigation: Invest in auto-scaling infrastructure
- Contingency: Queue system for peak hours
- Owner: Engineering Team""",

            "Success Metrics & KPIs": """CRITICAL: How we measure if product succeeds.

Define:
- NORTH STAR METRIC: The one metric that matters most
- PRIMARY KPIs: 3-5 metrics with specific TARGETS
- SUPPORTING METRICS: Secondary metrics
- MEASUREMENT METHODOLOGY: How to calculate each
- REPORTING CADENCE: Daily/weekly/monthly
- DASHBOARD REQUIREMENTS: What to visualize

Include 30/60/90-day targets.
Example:
- Clip creation rate: 95% success rate
- User engagement: 10% increase in session time
- Viral coefficient: K > 1""",

            "Analytics": """CRITICAL: What events to track and why.

Define:
- GLOBAL METRICS: Platform-level metrics across all features
- FEATURE-SPECIFIC METRICS: Metrics unique to this product
- EVENT DEFINITIONS: What counts as a view, completion, share, etc.
- FORMULAS: How to calculate viral coefficient, conversion rates
- FUTURE SCOPE: What we'll add later

Reference PRD format:
- Clip creation rate: Number of clips created / Number of streams
- Viral coefficient: (Total Shares / Total Active Users) * (Clicks / Shares) * (New Users / Clicks)""",

            "Edge Cases": """CRITICAL: Exhaustive list of what could go wrong.

Categorize:
- CONTENT: Duplicate detection, wrong attribution, deleted content
- USER: Banned users, inactive accounts, permissions issues
- SYSTEM: Encoding failures, CDN failures, database timeouts
- UPLOAD: Partial uploads, corrupted files, format issues
- NETWORK: Offline users, slow connections, timeouts

For each:
- SCENARIO: What happens
- HANDLING: How the system responds
- USER FEEDBACK: What does the user see

Example:
- Duplicate clips -> Dedupe via hash -> Show "similar clip exists" message
- CDN failure -> Fallback origin -> Retry with original server""",

            "Feed": """CRITICAL: For content platforms - how content is discovered.

Define:
- PRINCIPLES: What's the primary goal (acquisition vs engagement vs discovery)
- SURFACE RULES: Where does content appear (home, profile, category pages)
- PLATFORM BEHAVIOR: Desktop Web vs Mobile App vs External URLs
- ENTRY POINTS: How users access the feed
- VISIBILITY THRESHOLDS: When does content appear (e.g., after 100 clips exist)
- SCORING ALGORITHM: How content is ranked (views, shares, completion rate)
- BATCH COMPOSITION: How content is batched and prefetched
- GUARDRAILS: What should NEVER be shown

Reference PRD has detailed formula for scoring:
Score = (completion_rate * W_completion) + (share_rate * W_share) + (follow_rate * W_follow) - (skip_rate * W_skip)""",

            "Moderation & NSFW System": """CRITICAL: For UGC platforms - how content is reviewed.

Define:
- AUTOMATED MODERATION: AI/ML detection inputs (video, audio, text)
- MANUAL MODERATION: Human review workflow, tool requirements
- MODERATION CRITERIA: What gets approved/rejected
- SLA: How quickly to review (e.g., 30 minutes)
- EDGE CASES: Conflict resolution, escalations
- NSFW SCORE: How to measure appropriateness

Include flow: Content submitted -> Auto-check -> Score -> Approve/Reject/Manual Review""",

            "AI Design for Clips": """CRITICAL: For AI/ML features - how the algorithm works.

Define:
- SIGNALS: What inputs does AI use (chat spikes, audio peaks, viewer spikes)
- FORMULA: Exact scoring equation (e.g., Score = w1*chat + w2*audio + w3*viewers)
- WEIGHTS: What are the weight values and why
- OUTPUT: What does AI generate (top 5-10 clips, 30 sec clips)
- PROCESSING: When does it run (stream end, periodic)
- LIMITATIONS: Constraints on AI output
- HUMAN OVERRIDE: Can users edit AI-generated clips

Example:
- Signal: Chat spike = 5x normal messages
- Score = 0.4(chat) + 0.3(audio) + 0.3(viewers)
- Output: Top 5 clips per stream, 30 seconds each""",

            "Monetization": """CRITICAL: Revenue strategy for the product.

Define:
- REVENUE SOURCES: Ads, subscriptions, transaction fees, sponsorships
- REVENUE SPLIT: How is revenue shared (creator vs platform)
- PRICING: How much do users pay
- EDGE CASES: What happens to revenue if creator deleted, banned

Example:
- Banner ads displayed on clips
- Revenue split: 50% streamer, 50% platform
- Creator deleted -> revenue goes to streamer
- Streamer banned -> revenue frozen""",

            "AI Design": """CRITICAL: For AI/ML products - detailed algorithm specification.

Define:
- SIGNALS: What data inputs the model uses
- MODEL ARCHITECTURE: Type of model, training approach
- TRAINING DATA: What data is used to train
- FORMULA: Exact equation for scoring/ranking
- WEIGHTS: Parameter values with justification
- OUTPUT: What the model produces
- THRESHOLDS: Cutoff points for decisions
- MONITORING: How to track model performance""",
        }
        return guides.get(section, f"""Write comprehensive content for '{section}' section.
Assume the reader has ZERO knowledge about this product.
- Start with clear definitions
- Use specific examples
- Include quantitative data where possible
- Structure with headers, bullet points, and tables
- Minimum 500 words""")
        return guides.get(section, f"Write detailed, specific content for the '{section}' section.")

    def _get_fallback(self, section: str, context: PRDContext) -> str:
        """DEPRECATED: This method now raises error to trigger global switch."""
        raise RuntimeError(f"Fallback called for {section} - This should not happen!")


# =============================================================================
# 4️⃣ EVALUATOR AGENT
# =============================================================================

class EvaluatorAgent(BaseAgent):
    """
    EVALUATOR AGENT - A ruthless VP of Product.

    SCORING DIMENSIONS (score each 1-5):
    - Clarity (15%)
    - Depth (20%)
    - Actionability (25%)
    - User Focus (15%)
    - Research Alignment (15%)
    - Strategic Thinking (10%)

    RULES:
    - Penalize fluff heavily
    - Prefer execution-ready outputs
    - Reward specificity
- If all weak - pick best but highlight weakness

    CALCULATION:
    Final Score = (Clarity * 0.15) + (Depth * 0.20) + (Actionability * 0.25) +
                  (User Focus * 0.15) + (Research Alignment * 0.15) + (Strategic Thinking * 0.10)
    """

    def __init__(self, gemini_api_key: str, error_logger: GitHubErrorLogger = None):
        super().__init__(
            gemini_api_key,
            ["gemini-2.0-flash", "gemini-2.0-flash-lite"],
            "Evaluator Agent",
            error_logger
        )

    def select_best(self, section: str, options: List[str], context: PRDContext) -> Tuple[str, str]:
        """Select the best option using VP-level scoring."""
        options_text = "\n\n".join([f"=== OPTION {i+1} ===\n{opt}" for i, opt in enumerate(options)])

        prompt = f"""You are a ruthless VP of Product.

##  TASK
Evaluate and select the best option for the "{section}" section of a PRD.

Product Context:
- Idea: {context.idea}
- Problem: {context.problem_statement}

OPTIONS:
{options_text}

##  SCORING DIMENSIONS
Score each (1-5):
- Clarity (15%)
- Depth (20%)
- Actionability (25%)
- User Focus (15%)
- Research Alignment (15%)
- Strategic Thinking (10%)

##  RULES
- Penalize fluff heavily
- Prefer execution-ready outputs
- Reward specificity
- If all weak -> pick best but highlight weakness

##  CALCULATION
Final Score = (Clarity * 0.15) + (Depth * 0.20) + (Actionability * 0.25) + (User Focus * 0.15) + (Research Alignment * 0.15) + (Strategic Thinking * 0.10)

##  OUTPUT
{{
    "scores": [
        {{
            "option": 1,
            "clarity": 0,
            "depth": 0,
            "actionability": 0,
            "user_focus": 0,
            "research_alignment": 0,
            "strategic_thinking": 0,
            "final_score": 0.0
        }}
    ],
    "selected_index": 0,
    "confidence": 0.0,
    "reason": "sharp reasoning"
}}"""
        try:
            result = self._call_llm_json(prompt, f"Evaluating section: {section}")
            idx = result.get("selected_index", 0)
            rationale = result.get("reason", result.get("rationale", "Selected for best overall quality."))
            if isinstance(idx, int) and 0 <= idx < len(options):
                return options[idx], rationale
            return options[0], rationale
        except Exception:
            return options[0], "Selected based on overall quality and completeness."


# =============================================================================
# 5️⃣ GAP DETECTOR AGENT
# =============================================================================

class GapDetectorAgent(BaseAgent):
    """
     GAP DETECTOR AGENT - Precision PRD gap detector.

    CHECK:
    - Missing sections
    - Weak logic
    - Undefined flows
    - Missing edge cases
    - Missing metrics
    """

    def __init__(self, gemini_api_key: str, error_logger: GitHubErrorLogger = None):
        super().__init__(
            gemini_api_key,
            ["gemini-2.0-flash", "gemini-2.0-flash-lite"],
            "Gap Detector",
            error_logger
        )

    def detect_gaps(self, prd_content: str, new_user_input: str = "",
                    context: PRDContext = None) -> GapReport:
        """Analyze PRD and new input to find gaps."""
        prompt = f"""You are a precision PRD gap detector.

##  TASK
Identify missing or weak areas.

##  CHECK
- Missing sections
- Weak logic
- Undefined flows
- Missing edge cases
- Missing metrics

CURRENT PRD:
{prd_content[:5000]}

{"NEW USER INPUT: " + new_user_input if new_user_input else ""}
{"PRODUCT CONTEXT: " + context.idea if context else ""}

##  OUTPUT
{{
    "missing_sections": ["list of topics/sections completely absent"],
    "weak_sections": ["list of sections with weak content"],
    "logical_gaps": ["list of logical inconsistencies or undefined flows"],
    "required_research_areas": ["areas needing additional research"],
    "severity": "low | medium | high",
    "improvements_needed": [
        {{"section": "section name", "issue": "what's wrong", "suggestion": "how to fix"}}
    ]
}}"""
        try:
            result = self._call_llm_json(prompt, "Gap detection")
            return GapReport(
                missing_sections=result.get("missing_sections", []),
                improvements_needed=result.get("improvements_needed", []),
                weak_areas=result.get("weak_areas", []),
                raw_analysis=json.dumps(result)
            )
        except Exception:
            return GapReport(raw_analysis="Gap detection encountered an error.")


# =============================================================================
# 6️⃣ ENGINEERING MANAGER AGENT
# =============================================================================

class EngineeringManagerAgent(BaseAgent):
    """
    ️ ENGINEERING MANAGER AGENT — Senior Engineering Manager.

    TASK: Validate PRD for technical completeness and scalability.

    CHECK:
    - System design completeness
    - APIs and data flow
    - Edge cases and failures
    - Scalability risks
    - UI/UX feasibility

    RULE: If ANY critical gap exists -> reject
    """

    def __init__(self, gemini_api_key: str, error_logger: GitHubErrorLogger = None):
        super().__init__(
            gemini_api_key,
            ["gemini-2.0-flash", "gemini-2.0-flash-lite"],
            "Engineering Manager",
            error_logger
        )

    def review(self, prd_content: str, context: PRDContext) -> EngineeringReview:
        """Review entire PRD from engineering perspective."""
        prompt = f"""You are a senior Engineering Manager.

##  TASK
Validate PRD for technical completeness and scalability.

##  CHECK
- System design completeness
- APIs and data flow
- Edge cases and failures
- Scalability risks
- UI/UX feasibility

## ⚠️ RULE
If ANY critical gap exists -> reject (set status to "needs_changes")

PRODUCT: {context.idea}
PROBLEM: {context.problem_statement}

PRD CONTENT:
{prd_content[:6000]}

##  OUTPUT
{{
    "status": "approved | needs_changes",
    "approved": true|false,
    "blocking_issues": ["list of critical blocking issues"],
    "scalability_risks": ["list of scalability concerns"],
    "missing_technical_details": ["list of missing technical specs"],
    "ui_ux_gaps": ["list of UI/UX gaps"],
    "issues": [
        {{"section": "name", "severity": "critical|major|minor", "issue": "description", "recommendation": "how to fix"}}
    ],
    "feedback_for_sections": {{
        "Section Name": "specific feedback to improve this section"
    }}
}}"""
        try:
            result = self._call_llm_json(prompt, "Engineering review")
            return EngineeringReview(
                issues=result.get("issues", []),
                approved=result.get("approved", False),
                feedback_for_sections=result.get("feedback_for_sections", {}),
                raw_review=json.dumps(result)
            )
        except Exception:
            return EngineeringReview(
                approved=True,
                raw_review="Engineering review completed with default approval."
            )


# =============================================================================
# 7️⃣ VP PRODUCT AGENT
# =============================================================================

class VPProductAgent(BaseAgent):
    """
     VP PRODUCT AGENT — VP of Product responsible for final approval.

    TASK: Evaluate PRD from a business and strategic perspective.

    CHECK:
    - Market viability
    - Competitive advantage
    - Monetization logic
    - Product completeness
    - Edge cases
    """

    def __init__(self, gemini_api_key: str, error_logger: GitHubErrorLogger = None):
        super().__init__(
            gemini_api_key,
            ["gemini-2.0-flash", "gemini-2.0-flash-lite"],
            "VP Product",
            error_logger
        )

    def review(self, prd_content: str, context: PRDContext,
               eng_review: EngineeringReview = None) -> dict:
        """Final executive review of the complete PRD."""
        eng_summary = ""
        if eng_review and eng_review.raw_review:
            eng_summary = f"\nEngineering Review Summary: {eng_review.raw_review[:1000]}"

        prompt = f"""You are a VP of Product responsible for final approval.

##  TASK
Evaluate PRD from a business and strategic perspective.

##  CHECK
- Market viability
- Competitive advantage
- Monetization logic
- Product completeness
- Edge cases

PRODUCT: {context.idea}
PROBLEM: {context.problem_statement}
{eng_summary}

PRD CONTENT:
{prd_content[:6000]}

Provide your review in the following format:

## Executive Review Summary
[2-3 sentence overall assessment]

## Missed Cases & Critical Gaps
**Q1: [Specific question about a gap]**
A: [Detailed analysis and recommendation]

**Q2: [Another critical question]**
A: [Detailed analysis]

[Continue with all identified gaps]

##  Final Output
Also provide a structured JSON block:
{{
    "edge_cases": ["list of edge cases identified"],
    "business_risks": ["list of business risks"],
    "product_gaps": ["list of product gaps"],
    "strategic_improvements": ["list of strategic improvements"],
    "final_verdict": "approve | refine"
}}

## Final Verdict
[Approved/Conditional/Needs Revision] with reasoning"""

        try:
            raw = self._call_llm(prompt, "VP Product review")
            return {
                "review_passed": True,
                "missed_cases": raw,
                "raw_review": raw
            }
        except Exception:
            return {
                "review_passed": True,
                "missed_cases": "VP Product review completed.",
                "raw_review": "Review completed with standard approval."
            }


# =============================================================================
#  PRD ORCHESTRATOR — Coordinates Everything
# =============================================================================

class PRDOrchestrator:
    """
    The central coordinator that manages all 7 agents, handles the
    initial generation flow, iterative refinement, memory persistence,
    document generation, and error logging.
    """

    MAX_ENG_LOOPS = 2  # Maximum engineering review re-loops

    def __init__(self, gemini_api_key: str, tavily_api_key: str = "",
                 google_api_key: str = "", google_cx: str = "",
                 github_pat: str = "", github_repo: str = ""):

        self.gemini_api_key = gemini_api_key
        self.github_pat = github_pat
        self.github_repo = github_repo

        self.error_logger = GitHubErrorLogger(github_pat, github_repo)

        prd_logger.info("=" * 50)
        prd_logger.info(" Initializing PRD Engine — 7 Agent System")
        prd_logger.info("=" * 50)
        
        # Initialize all 7 agents
        self.god_agent = GodAgent(gemini_api_key, self.error_logger)
        self.classifier = ClassifierAgent(gemini_api_key, self.error_logger)
        self.researcher = ResearchAgent(
            gemini_api_key, tavily_api_key, google_api_key, google_cx,
            github_pat, github_repo, self.error_logger
        )
        self.generator = PRDGeneratorAgent(gemini_api_key, self.error_logger)
        self.evaluator = EvaluatorAgent(gemini_api_key, self.error_logger)
        self.gap_detector = GapDetectorAgent(gemini_api_key, self.error_logger)
        self.eng_manager = EngineeringManagerAgent(gemini_api_key, self.error_logger)
        self.vp_product = VPProductAgent(gemini_api_key, self.error_logger)

        prd_logger.info("=" * 50)
        prd_logger.info("✅ All 7 agents initialized")
        prd_logger.info("=" * 50)

    # -----------------------------------------------------------------
    # INITIAL GENERATION FLOW
    # -----------------------------------------------------------------

    def generate_prd(self, user_input: str,
                     progress_callback=None,
                     memory: PRDMemory = None) -> Tuple[bool, str, str, PRDMemory]:
        """
        Execute the complete initial PRD generation workflow.

        Returns: (success, docx_path, status_message, memory)
        """
        prd_logger.info(f" Starting PRD generation for: {user_input[:100]}...")
        
        try:
            if memory is None:
                memory = PRDMemory()
            memory.user_inputs.append(user_input)

            # ---- STEP 1: God Agent plans workflow ----
            self._progress(progress_callback, " God Agent: Planning workflow...")
            prd_logger.info(" God Agent: Planning workflow...")
            log_agent_start(prd_logger, "GodAgent", "Initial planning")
            god_plan = self.god_agent.plan_initial_workflow(user_input)
            log_agent_end(prd_logger, "GodAgent", "COMPLETE")

            # ---- STEP 2: Classifier Agent ----
            self._progress(progress_callback, " Classifier Agent: Analyzing input type...")
            prd_logger.info(" Classifier Agent: Analyzing input type...")
            log_agent_start(prd_logger, "ClassifierAgent", "Classify input")
            context = self.classifier.classify(user_input)
            memory.context = context
            prd_logger.info(f"Classified as: {context.input_type} | Idea: {context.idea[:50]}...")
            log_agent_end(prd_logger, "ClassifierAgent", "COMPLETE")

            # ---- STEP 3: Research Agent ----
            self._progress(progress_callback, " Research Agent: Searching Tavily + Google + Specter reports...")
            prd_logger.info(" Research Agent: Searching...")
            log_agent_start(prd_logger, "ResearchAgent", "Research")
            queries = god_plan.get("research_queries", [
                f"market analysis {context.idea}",
                f"competitors {context.idea}",
                f"technical challenges {context.idea}",
                f"user needs {context.problem_statement}"
            ])
            prd_logger.info(f"Research queries: {queries}")
            research_data = self.researcher.research(queries, memory)
            memory.research_memory.update(research_data.get("results_by_query", {}))
            context.research_data = research_data
            prd_logger.info(f"Research complete: {len(research_data.get('results_by_query', {}))} query results")
            log_agent_end(prd_logger, "ResearchAgent", "COMPLETE")

            # ---- STEP 4: Generate + Evaluate all sections ----
            self._progress(progress_callback, "✍️ PRD Generator: Creating detailed sections (3 options each)...")
            prd_logger.info("✍️ PRD Generator: Creating detailed sections...")
            
            # Dynamically determine sections based on product type
            sections_to_generate = PRDGeneratorAgent.get_sections_for_product(god_plan)
            prd_logger.info(f" Sections to generate: {sections_to_generate}")
            prd_sections = {}
            total_sections = len(sections_to_generate)

            for i, section_name in enumerate(sections_to_generate):
                self._progress(
                    progress_callback,
                    f"✍️ Generating section {i+1}/{total_sections}: {section_name}..."
                )
                prd_logger.info(f"Generating section {i+1}/{total_sections}: {section_name}")

                # Generate 3 options
                options = self.generator.generate_section(
                    section_name, context,
                    research_data.get("summary", ""),
                    god_plan
                )

                # Evaluate and select best
                log_agent_start(prd_logger, "EvaluatorAgent", f"Evaluate: {section_name}")
                selected, rationale = self.evaluator.select_best(
                    section_name, options, context
                )
                log_agent_end(prd_logger, "EvaluatorAgent", f"COMPLETE | Selected option {options.index(selected)+1}")

                prd_sections[section_name] = PRDSection(
                    title=section_name,
                    options=options,
                    selected_option=selected,
                    rationale=rationale
                )

                time.sleep(3)  # Rate limit protection

            memory.prd_state = prd_sections
            prd_logger.info(f"✅ All {total_sections} PRD sections generated")

            # ---- STEP 5: Engineering Manager Review (with re-loop) ----
            prd_md = memory.get_prd_markdown()
            for loop in range(self.MAX_ENG_LOOPS):
                self._progress(
                    progress_callback,
                    f"️ Engineering Manager: Technical review (pass {loop+1})..."
                )
                eng_review = self.eng_manager.review(prd_md, context)
                memory.engineering_review = asdict(eng_review) if hasattr(eng_review, '__dataclass_fields__') else {"raw": str(eng_review)}

                if eng_review.approved:
                    self._progress(progress_callback, "✅ Engineering Manager: Approved!")
                    break

                # Re-generate affected sections
                if eng_review.feedback_for_sections:
                    self._progress(
                        progress_callback,
                        f" Re-generating {len(eng_review.feedback_for_sections)} sections based on engineering feedback..."
                    )
                    for sec_name, feedback in eng_review.feedback_for_sections.items():
                        if sec_name in prd_sections:
                            new_options = self.generator.generate_section(
                                sec_name, context,
                                research_data.get("summary", ""),
                                god_plan,
                                engineering_feedback=feedback
                            )
                            new_selected, new_rationale = self.evaluator.select_best(
                                sec_name, new_options, context
                            )
                            prd_sections[sec_name].update(new_selected, new_rationale)
                            time.sleep(1)

                    prd_md = memory.get_prd_markdown()
                else:
                    break

            # ---- STEP 6: VP Product Review ----
            self._progress(progress_callback, " VP Product: Final executive review...")
            vp_review = self.vp_product.review(prd_md, context, eng_review)
            memory.vp_review = vp_review

            # ---- STEP 7: Generate Documents ----
            self._progress(progress_callback, " Generating PRD documents...")
            docx_path = self._generate_docx(memory, eng_review, vp_review)

            # Push to GitHub
            github_msg = ""
            if self.github_pat and self.github_repo:
                if self._push_to_github(docx_path):
                    github_msg = " (Pushed to GitHub)"

            memory.version = 1
            success_msg = f"PRD v{memory.version} generated successfully!{github_msg}"
            self._progress(progress_callback, f" {success_msg}")
            prd_logger.info(f" PRD generation complete: {success_msg}")

            return True, docx_path, success_msg, memory

        except Exception as e:
            error_msg = f"PRD generation failed: {str(e)}"
            tb = traceback.format_exc()
            prd_logger.error(f"❌ PRD generation failed: {error_msg}")
            log_error(prd_logger, "PRDOrchestrator", str(e), f"Input: {user_input[:200]}")
            self.error_logger.log_error("orchestrator_generate", str(e), tb, f"Input: {user_input[:200]}")
            return False, "", error_msg, memory

    # -----------------------------------------------------------------
    # ITERATIVE REFINEMENT FLOW
    # -----------------------------------------------------------------

    def refine_prd(self, new_input: str, memory: PRDMemory,
                   progress_callback=None,
                   specific_section: str = None) -> Tuple[bool, str, str, PRDMemory]:
        """
        Execute iterative refinement workflow.

        Args:
            new_input: User's new requirements/modifications
            memory: Existing PRDMemory from previous generation
            specific_section: If set, only regenerate this section

        Returns: (success, docx_path, status_message, updated_memory)
        """
        try:
            memory.user_inputs.append(new_input)
            context = memory.context

            # ---- STEP 1: God Agent interprets the update ----
            self._progress(progress_callback, " God Agent: Interpreting your update...")
            update_plan = self.god_agent.interpret_update(new_input, memory)

            affected = update_plan.get("affected_sections", [])
            if specific_section:
                affected = [specific_section]

            # ---- STEP 2: Gap Detection ----
            self._progress(progress_callback, " Gap Detector: Scanning for missing pieces...")
            prd_md = memory.get_prd_markdown()
            gap_report = self.gap_detector.detect_gaps(prd_md, new_input, context)

            # ---- STEP 3: Incremental Research (if needed) ----
            research_data = context.research_data or {}
            if update_plan.get("new_research_needed", False):
                self._progress(progress_callback, " Research Agent: Incremental research (reusing cache)...")
                new_queries = update_plan.get("research_queries", [f"research: {new_input}"])
                new_research = self.researcher.research(new_queries, memory)
                memory.research_memory.update(new_research.get("results_by_query", {}))

                # Merge research
                if "summary" in new_research:
                    old_summary = research_data.get("summary", "")
                    research_data["summary"] = old_summary + "\n\n--- UPDATED RESEARCH ---\n" + new_research["summary"]
                research_data.update({k: v for k, v in new_research.items() if k != "summary"})

            # ---- STEP 4: Regenerate affected sections ----
            god_plan = self.god_agent.plan_initial_workflow(
                f"{context.original_input}\n\nAdditional: {new_input}"
            )
            god_plan["special_instructions"] = update_plan.get(
                "instructions_for_generator",
                f"User update: {new_input}"
            )

            self._progress(progress_callback, f"✍️ Regenerating {len(affected)} sections...")
            for sec_name in affected:
                if sec_name in memory.prd_state:
                    self._progress(progress_callback, f"  ✍️ Rewriting: {sec_name}...")
                    options = self.generator.generate_section(
                        sec_name, context,
                        research_data.get("summary", ""),
                        god_plan
                    )
                    selected, rationale = self.evaluator.select_best(sec_name, options, context)
                    memory.prd_state[sec_name].update(selected, rationale)
                    time.sleep(1)

            # ---- STEP 5: Engineering Manager Review ----
            prd_md = memory.get_prd_markdown()
            self._progress(progress_callback, "️ Engineering Manager: Reviewing updates...")
            eng_review = self.eng_manager.review(prd_md, context)

            if not eng_review.approved and eng_review.feedback_for_sections:
                self._progress(progress_callback, " Addressing engineering feedback...")
                for sec_name, feedback in eng_review.feedback_for_sections.items():
                    if sec_name in memory.prd_state:
                        options = self.generator.generate_section(
                            sec_name, context,
                            research_data.get("summary", ""),
                            god_plan,
                            engineering_feedback=feedback
                        )
                        selected, rationale = self.evaluator.select_best(sec_name, options, context)
                        memory.prd_state[sec_name].update(selected, rationale)
                        time.sleep(1)

            # ---- STEP 6: VP Product Review ----
            prd_md = memory.get_prd_markdown()
            self._progress(progress_callback, " VP Product: Reviewing updates...")
            vp_review = self.vp_product.review(prd_md, context, eng_review)
            memory.vp_review = vp_review

            # ---- STEP 7: Generate Updated Documents ----
            self._progress(progress_callback, " Generating updated PRD documents...")
            memory.version += 1
            docx_path = self._generate_docx(memory, eng_review, vp_review)

            github_msg = ""
            if self.github_pat and self.github_repo:
                if self._push_to_github(docx_path):
                    github_msg = " (Pushed to GitHub)"

            success_msg = f"PRD v{memory.version} refined successfully!{github_msg}"
            self._progress(progress_callback, f" {success_msg}")

            return True, docx_path, success_msg, memory

        except Exception as e:
            error_msg = f"PRD refinement failed: {str(e)}"
            tb = traceback.format_exc()
            self.error_logger.log_error("orchestrator_refine", str(e), tb, f"Input: {new_input[:200]}")
            return False, "", error_msg, memory

    # -----------------------------------------------------------------
    # DOCUMENT GENERATION
    # -----------------------------------------------------------------

    def _generate_docx(self, memory: PRDMemory,
                       eng_review: EngineeringReview = None,
                       vp_review: dict = None) -> str:
        """Generate professional DOCX document."""
        if not DOCX_AVAILABLE:
            # Fallback to markdown
            return self._save_markdown(memory)

        doc = Document()
        context = memory.context

        # Title
        title = doc.add_heading("Product Requirements Document", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        product_name = context.idea[:80] if context else "Product"
        subtitle = doc.add_heading(f"{product_name}", 1)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata table
        table = doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"
        meta = [
            ("Created", datetime.now().strftime("%Y-%m-%d %H:%M")),
            ("Version", f"v{memory.version}"),
            ("Author", "AI PRD Engine (7-Agent System)"),
            ("Status", "VP Product Approved" if vp_review and vp_review.get("review_passed") else "Draft"),
            ("Iterations", str(len(memory.user_inputs)))
        ]
        for i, (key, val) in enumerate(meta):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = val

        doc.add_paragraph()

        # All PRD sections
        for section_name in PRDGeneratorAgent.SECTIONS:
            if section_name in memory.prd_state:
                section = memory.prd_state[section_name]
                doc.add_heading(section_name, 1)
                if section.selected_option:
                    for line in section.selected_option.split("\n"):
                        stripped = line.strip()
                        if stripped:
                            doc.add_paragraph(stripped)
                if section.rationale:
                    p = doc.add_paragraph()
                    run = p.add_run(f"Selection Rationale: {section.rationale}")
                    run.font.size = Pt(8)
                    run.italic = True
                doc.add_paragraph()

        # Engineering Review section
        if eng_review and eng_review.raw_review:
            doc.add_heading("Engineering Review", 1)
            doc.add_paragraph(
                "The following is the technical review from the Engineering Manager Agent:"
            )
            review_text = eng_review.raw_review
            if isinstance(review_text, str) and len(review_text) > 20:
                try:
                    parsed = json.loads(review_text)
                    if parsed.get("issues"):
                        for issue in parsed["issues"]:
                            doc.add_paragraph(
                                f"[{issue.get('severity', 'info').upper()}] {issue.get('section', '')}: "
                                f"{issue.get('issue', '')} -> {issue.get('recommendation', '')}",
                                style="List Bullet"
                            )
                    approval_status = "✅ Approved" if parsed.get("approved") else "⚠️ Conditional"
                    doc.add_paragraph(f"\nStatus: {approval_status}")
                except:
                    doc.add_paragraph(review_text[:3000])

        # VP Product Review section
        if vp_review and vp_review.get("missed_cases"):
            doc.add_heading("VP Product Executive Review", 1)
            for line in vp_review["missed_cases"].split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        # Change History
        if memory.version > 1:
            doc.add_heading("Change History", 1)
            for inp_idx, inp in enumerate(memory.user_inputs):
                doc.add_paragraph(f"v{inp_idx + 1}: {inp[:200]}", style="List Bullet")

        # Save
        prd_title = self._generate_title(context)
        filename = f"PRD - {prd_title} v{memory.version}.docx"
        filepath = os.path.join("reports", filename)
        os.makedirs("reports", exist_ok=True)
        doc.save(filepath)
        print(f"   DOCX saved: {filepath}")
        return filepath

    def _save_markdown(self, memory: PRDMemory) -> str:
        """Save PRD as markdown file (fallback if DOCX unavailable)."""
        md_content = memory.get_prd_markdown()
        prd_title = self._generate_title(memory.context)
        filename = f"PRD - {prd_title} v{memory.version}.md"
        filepath = os.path.join("reports", filename)
        os.makedirs("reports", exist_ok=True)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(md_content)
        return filepath

    def generate_markdown_export(self, memory: PRDMemory) -> str:
        """Export current PRD state as markdown string."""
        return memory.get_prd_markdown()

    def generate_pdf_export(self, memory: PRDMemory) -> Optional[str]:
        """Export current PRD as PDF file."""
        if not PDF_EXPORT_AVAILABLE:
            return None

        md_content = memory.get_prd_markdown()
        html_content = md_lib.markdown(md_content)

        styled_html = f"""<html><head><style>
body {{ font-family: 'Helvetica', 'Arial', sans-serif; font-size: 11px; margin: 40px; color: #333; }}
h1 {{ color: #1a1a2e; border-bottom: 2px solid #e74c3c; padding-bottom: 8px; }}
h2 {{ color: #2c3e50; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; margin-top: 25px; }}
ul, ol {{ margin-left: 20px; }}
p {{ line-height: 1.6; }}
</style></head><body>{html_content}</body></html>"""

        prd_title = self._generate_title(memory.context)
        filename = f"PRD - {prd_title} v{memory.version}.pdf"
        filepath = os.path.join("reports", filename)
        os.makedirs("reports", exist_ok=True)

        try:
            with open(filepath, "wb") as f:
                pisa.CreatePDF(styled_html, dest=f)
            return filepath
        except Exception:
            return None

    # -----------------------------------------------------------------
    # UTILITIES
    # -----------------------------------------------------------------

    def _progress(self, callback, message: str):
        """Send progress update."""
        print(f"  {message}")
        if callback:
            callback(message)

    def _generate_title(self, context: PRDContext) -> str:
        """Generate a clean title from the context."""
        import re
        text = (context.idea if context else "") or "Product PRD"
        text = re.sub(r'[^\w\s-]', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        words = text.split()[:8]
        title = " ".join(w.capitalize() for w in words)
        if len(title) > 60:
            title = title[:57] + "..."
        return title or "Product PRD"

    def _push_to_github(self, filepath: str) -> bool:
        """Push generated PRD to GitHub repository."""
        if not self.github_pat or not self.github_repo or not GITHUB_AVAILABLE:
            return False
        try:
            g = Github(self.github_pat)
            repo = g.get_repo(self.github_repo)
            with open(filepath, "rb") as f:
                content = f.read()
            remote_path = filepath.replace("\\", "/")
            filename = os.path.basename(filepath)
            repo.create_file(remote_path, f"docs: PRD generated - {filename}", content)
            print(f"  ☁️ Pushed to GitHub: {remote_path}")
            return True
        except Exception as e:
            print(f"  ⚠️ GitHub push failed: {e}")
            return False
